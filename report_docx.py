"""
Geração do Relatório Gerencial da Pesquisa de Preços em .docx.

Consome as estruturas já calculadas pelo app (resumo do processo, cobertura de
itens e categorização), sem refazer nenhum cálculo — este módulo apenas
formata. Referências normativas: Lei nº 14.133/2021 e IN SEGES/ME nº 65/2021.
"""
from __future__ import annotations

import io
from datetime import datetime, timezone

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _fmt_data_iso(iso: str | None) -> str:
    if not iso:
        return "—"
    try:
        d = datetime.fromisoformat(iso)
        return d.strftime("%d/%m/%Y %H:%M")
    except Exception:
        return str(iso)


def _delta_horas_txt(iso_inicio: str | None, iso_fim: str | None) -> str:
    if not iso_inicio or not iso_fim:
        return "—"
    try:
        a = datetime.fromisoformat(iso_inicio)
        b = datetime.fromisoformat(iso_fim)
        if a.tzinfo is not None:
            a = a.astimezone(timezone.utc).replace(tzinfo=None)
        if b.tzinfo is not None:
            b = b.astimezone(timezone.utc).replace(tzinfo=None)
        h = (b - a).total_seconds() / 3600
        if h < 0:
            return "—"
        return f"{h:.1f}h ({h / 24.0:.1f} dias)"
    except Exception:
        return "—"


def _status_participacao(p: dict) -> str:
    if p.get("enviou_orcamento"):
        return "Enviou orçamento"
    if p.get("recusou"):
        return "Recusou"
    if p.get("fez_pergunta"):
        return "Fez pergunta"
    if p.get("confirmou_leitura"):
        return "Confirmou leitura"
    return "Sem resposta"


def _add_heading(doc: Document, texto: str, nivel: int = 1) -> None:
    doc.add_heading(texto, level=nivel)


def _add_tabela(doc: Document, cabecalho: list[str], linhas: list[list[str]]) -> None:
    tabela = doc.add_table(rows=1, cols=len(cabecalho))
    tabela.style = "Light Grid Accent 1"
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tabela.rows[0].cells
    for i, titulo in enumerate(cabecalho):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(titulo)
        run.bold = True
    for linha in linhas:
        cells = tabela.add_row().cells
        for i, valor in enumerate(linha):
            cells[i].text = str(valor if valor is not None else "—")
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Relatório
# ---------------------------------------------------------------------------

def build_report_docx(
    processo: dict,
    resumo: dict,
    cobertura_resumo: dict,
    categorias: dict[int, list[tuple[str, str]]] | None = None,
    rotulo_tipo_fn=None,
) -> bytes:
    """
    Gera o Relatório Gerencial da Pesquisa de Preços e retorna os bytes do .docx.

    processo:         dict com 'numero' e 'titulo'
    resumo:           saída de process_db.get_resumo_processo (enriquecida pelo app)
    cobertura_resumo: saída de _resumo_cobertura_itens
    categorias:       saída de _categorizar_itens_por_orcamentos (chaves 0, 1, 3)
    rotulo_tipo_fn:   função opcional para rotular tipos de e-mail
    """
    doc = Document()

    estilo = doc.styles["Normal"]
    estilo.font.name = "Calibri"
    estilo.font.size = Pt(11)

    # --- Cabeçalho -----------------------------------------------------------
    titulo = doc.add_heading("Relatório Gerencial da Pesquisa de Preços", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Processo: {processo.get('numero', '—')}").bold = True
    if processo.get("titulo"):
        doc.add_paragraph(processo["titulo"]).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        f"Gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M')}"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # --- 1. Síntese ----------------------------------------------------------
    _add_heading(doc, "1. Síntese da pesquisa")
    tempo_medio = resumo.get("tempo_medio_resposta_h")
    tempo_medio_txt = (
        f"{float(tempo_medio):.1f}h ({float(tempo_medio) / 24.0:.1f} dias)"
        if tempo_medio is not None else "não aferível com os dados disponíveis"
    )
    _add_tabela(
        doc,
        ["Indicador", "Valor"],
        [
            ["E-mails processados", resumo.get("total_emails", 0)],
            ["Fornecedores identificados", resumo.get("total_fornecedores", 0)],
            ["Fornecedores que enviaram orçamento", resumo.get("enviaram_orcamento", 0)],
            ["Fornecedores que recusaram", resumo.get("recusaram", 0)],
            ["Fornecedores sem resposta", resumo.get("sem_resposta", 0)],
            ["Tempo médio de resposta", tempo_medio_txt],
        ],
    )

    # --- 2. Classificação dos e-mails ---------------------------------------
    _add_heading(doc, "2. Classificação dos e-mails")
    por_tipo = resumo.get("por_tipo") or {}
    if por_tipo:
        linhas_tipo = []
        for tipo, qtd in sorted(por_tipo.items(), key=lambda x: -x[1]):
            rotulo = rotulo_tipo_fn(tipo) if rotulo_tipo_fn else tipo
            # remove emojis de rótulo para documento formal
            rotulo = "".join(ch for ch in str(rotulo) if ord(ch) < 0x2190).strip()
            linhas_tipo.append([rotulo or tipo, qtd])
        _add_tabela(doc, ["Categoria", "Quantidade"], linhas_tipo)
    else:
        doc.add_paragraph("Nenhum e-mail classificado neste processo.")

    # --- 3. Participação dos fornecedores ------------------------------------
    _add_heading(doc, "3. Participação dos fornecedores")
    participacoes = resumo.get("participacoes") or []
    itens_orcados_map = resumo.get("itens_orcados_por_fornecedor") or {}
    if participacoes:
        linhas_part = []
        for p_ in participacoes:
            email_f = (p_.get("email") or "").lower()
            linhas_part.append([
                p_.get("nome") or email_f or "—",
                _status_participacao(p_),
                _fmt_data_iso(p_.get("data_pedido_enviado")),
                _fmt_data_iso(p_.get("data_primeira_resposta")),
                _delta_horas_txt(p_.get("data_pedido_enviado"), p_.get("data_primeira_resposta")),
                itens_orcados_map.get(email_f, 0) or "—",
            ])
        _add_tabela(
            doc,
            ["Fornecedor", "Status", "Pedido enviado", "1ª resposta", "Tempo de resposta", "Itens orçados"],
            linhas_part,
        )
    else:
        doc.add_paragraph("Nenhuma participação registrada.")

    # --- 4. Cobertura de itens ------------------------------------------------
    _add_heading(doc, "4. Cobertura de orçamentos por item")
    total_itens = int(cobertura_resumo.get("total_itens", 0))
    if total_itens:
        _add_tabela(
            doc,
            ["Faixa", "Itens", "% do total"],
            [
                ["Zero orçamentos", int(cobertura_resumo.get("itens_zero", 0)),
                 f"{float(cobertura_resumo.get('perc_zero', 0)):.1f}%"],
                ["1 a 2 orçamentos", int(cobertura_resumo.get("itens_1_a_2", 0)),
                 f"{float(cobertura_resumo.get('perc_1_a_2', 0)):.1f}%"],
                ["3 ou mais orçamentos", int(cobertura_resumo.get("itens_3_ou_mais", 0)),
                 f"{float(cobertura_resumo.get('perc_3_ou_mais', 0)):.1f}%"],
            ],
        )
        if categorias:
            for chave, titulo_cat in ((0, "Itens sem nenhum orçamento"),
                                       (1, "Itens com 1 a 2 orçamentos"),
                                       (3, "Itens com 3 ou mais orçamentos")):
                itens_cat = categorias.get(chave) or []
                if itens_cat:
                    _add_heading(doc, titulo_cat, nivel=2)
                    for num, desc in sorted(itens_cat, key=lambda x: (str(x[0]), str(x[1]))):
                        doc.add_paragraph(f"{num} — {desc}", style="List Bullet")
    else:
        doc.add_paragraph("Não há itens consolidados para este processo.")

    # --- 5. Análise técnica ----------------------------------------------------
    _add_heading(doc, "5. Análise técnica")
    n_orc = int(resumo.get("enviaram_orcamento", 0))
    n_forn = int(resumo.get("total_fornecedores", 0))
    itens_criticos = int(cobertura_resumo.get("itens_zero", 0)) + int(cobertura_resumo.get("itens_1_a_2", 0))
    doc.add_paragraph(
        f"A pesquisa de preços consolidou respostas de {n_orc} fornecedor(es), de um universo de "
        f"{n_forn} fornecedor(es) identificado(s) nos e-mails do processo. "
        f"Dos {total_itens} item(ns) do processo, {int(cobertura_resumo.get('itens_3_ou_mais', 0))} "
        f"atingiram a referência desejável de 3 ou mais orçamentos, e {itens_criticos} item(ns) "
        f"permanecem abaixo dessa referência."
    )
    if itens_criticos:
        doc.add_paragraph(
            "Recomenda-se nova rodada de cotação direcionada aos itens com cobertura insuficiente, "
            "priorizando fornecedores que já orçaram itens de mesma natureza, ou a complementação da "
            "cesta de preços com outras fontes admitidas pelo art. 5º da IN SEGES/ME nº 65/2021 "
            "(painéis de preços, contratações similares, mídia especializada e sítios eletrônicos "
            "especializados), em atenção ao art. 23 da Lei nº 14.133/2021."
        )
    doc.add_paragraph(
        "Os preços consolidados no Mapa Comparativo devem ser avaliados criticamente quanto a "
        "valores inexequíveis ou excessivos, nos termos do art. 6º da IN SEGES/ME nº 65/2021. "
        "Os itens sinalizados na aba \"Revisar Casamentos\" e os alertas de preço fora da curva "
        "(método IQR) devem ser conferidos manualmente antes da adoção do menor preço como referência."
    )

    # --- 6. Limitações metodológicas -------------------------------------------
    _add_heading(doc, "6. Limitações metodológicas")
    doc.add_paragraph(
        "O pacote de e-mails normalmente não contém a lista completa de fornecedores convidados "
        "(envios em massa/cópia oculta). Assim, a taxa de resposta é calculada sobre os fornecedores "
        "que interagiram no processo, e não sobre o universo total de convidados.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Quando o pedido de cotação original não consta como e-mail isolado no lote, a data do pedido "
        "é inferida a partir do histórico citado no corpo das respostas, o que pode introduzir "
        "imprecisão no cálculo do tempo de resposta.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Itens sem valor preenchido na planilha de um fornecedor são tratados como \"não orçados por "
        "aquele fornecedor\", e não como falha de extração.",
        style="List Bullet",
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
