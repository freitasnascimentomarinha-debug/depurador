"""
Extração de texto de PDF (com OCR para escaneados), Word e Excel,
e interpretação estruturada via LLM (OpenRouter).
"""
from __future__ import annotations

import io
import json
import os
import re
import shutil
import time
from collections import Counter

import requests

from confidence import calcular_confianca_estrutural
from normalize_utils import (
    extrair_cnpj,
    extrair_razao_social,
    limpar_quebras_e_caracteres,
    normalizar_item,
)
from structured_extract import (
    detectar_tipo_e_rotear,
    extrair_docx_estruturado,
    extrair_xlsx_estruturado,
    tentar_extracao_estrutural_pdf,
)
from text_similarity import token_set_ratio

BOILERPLATE_MAX_LEN_SEM_DIGITO = 150
UNIT_TOKENS = (
    "UN", "UND", "UNID", "UNIDADE", "PCT", "PC", "PÇ", "EMB", "KG", "G", "MG",
    "L", "LT", "ML", "CX", "FR", "FD", "RL", "M", "M2", "M3", "PAR", "CJ",
)

# Precos aproximados por 1M tokens (USD) para estimativa de custo quando a API
# nao retornar custo total explicito.
MODEL_PRICING_PER_MILLION = {
    "openai/gpt-4o-mini": {"input": 0.15, "output": 0.60},
    "anthropic/claude-3-haiku": {"input": 0.25, "output": 1.25},
}

_OCR_MIN_ALNUM = 25
_OCR_MIN_TOKENS = 6


def _to_int(value) -> int:
    try:
        return int(value)
    except (TypeError, ValueError):
        return 0


def _extract_usage_info(response_json: dict, model: str) -> dict:
    usage = response_json.get("usage") or {}
    prompt_tokens = _to_int(usage.get("prompt_tokens") or usage.get("input_tokens"))
    completion_tokens = _to_int(usage.get("completion_tokens") or usage.get("output_tokens"))
    total_tokens = _to_int(usage.get("total_tokens"))
    if total_tokens <= 0:
        total_tokens = prompt_tokens + completion_tokens

    cost_usd = usage.get("cost")
    if cost_usd is None:
        cost_usd = usage.get("total_cost")

    estimated = False
    if cost_usd is None:
        pricing = MODEL_PRICING_PER_MILLION.get(model)
        if pricing:
            cost_usd = (
                (prompt_tokens / 1_000_000.0) * pricing["input"]
                + (completion_tokens / 1_000_000.0) * pricing["output"]
            )
            estimated = True
        else:
            cost_usd = 0.0
            estimated = True

    try:
        cost_usd = float(cost_usd)
    except (TypeError, ValueError):
        cost_usd = 0.0

    return {
        "prompt_tokens": prompt_tokens,
        "completion_tokens": completion_tokens,
        "total_tokens": total_tokens,
        "cost_usd": cost_usd,
        "estimated": estimated,
    }


def _import_fitz():
    raise RuntimeError("PyMuPDF desativado neste build")


def _import_pytesseract():
    import pytesseract
    return pytesseract


def _import_image():
    from PIL import Image
    return Image


def _import_document():
    from docx import Document
    return Document


def _import_openpyxl():
    import openpyxl
    return openpyxl


def _import_pdfplumber():
    import pdfplumber
    return pdfplumber


def _ocr_lang_candidates(ocr_lang: str) -> list[str]:
    base = (ocr_lang or "por").strip()
    candidates = [base]
    if "+" not in base:
        candidates.append(f"{base}+eng")
    if "eng" not in candidates:
        candidates.append("eng")
    vistos = set()
    unicos = []
    for cand in candidates:
        if cand and cand not in vistos:
            vistos.add(cand)
            unicos.append(cand)
    return unicos


def ocr_runtime_status() -> tuple[bool, str | None]:
    """Valida se o OCR esta operacional no ambiente atual."""
    if shutil.which("tesseract") is None:
        return False, "Binario 'tesseract' nao encontrado no PATH."

    try:
        pytesseract = _import_pytesseract()
        _ = pytesseract.get_tesseract_version()
    except Exception as exc:
        return False, f"Falha ao inicializar pytesseract/tesseract: {exc}"

    return True, None


def _texto_pobre_para_ocr(texto: str) -> bool:
    """Indica se o texto extraido da pagina esta fraco e deve tentar OCR."""
    base = re.sub(r"\s+", " ", str(texto or "")).strip()
    if not base:
        return True

    alnum = sum(1 for ch in base if ch.isalnum())
    tokens = re.findall(r"\w+", base)
    if alnum < _OCR_MIN_ALNUM:
        return True
    if len(tokens) < _OCR_MIN_TOKENS:
        return True
    return False


def _preferir_texto_ocr(texto_pdf: str, texto_ocr: str) -> bool:
    """Troca para OCR quando ele traz ganho claro de conteudo util."""
    base_pdf = re.sub(r"\s+", " ", str(texto_pdf or "")).strip()
    base_ocr = re.sub(r"\s+", " ", str(texto_ocr or "")).strip()
    if not base_ocr:
        return False
    if not base_pdf:
        return True
    return len(base_ocr) >= int(len(base_pdf) * 1.2)


def _ocr_image(image_obj, ocr_lang: str = "por") -> tuple[str, str | None]:
    ok_ocr, erro_ocr = ocr_runtime_status()
    if not ok_ocr:
        return "", erro_ocr

    pytesseract = _import_pytesseract()
    last_error = None
    for lang in _ocr_lang_candidates(ocr_lang):
        try:
            texto = pytesseract.image_to_string(image_obj, lang=lang)
            if texto and texto.strip():
                return texto.strip(), None
        except Exception as exc:
            last_error = str(exc)
    return "", last_error


def extract_text_from_image(path: str, ocr_lang: str = 'por') -> tuple[str, str | None]:
    Image = _import_image()
    with Image.open(path) as img:
        return _ocr_image(img, ocr_lang=ocr_lang)


# ---------- Extração de texto bruto ----------

def extract_text_from_pdf(path: str, ocr_lang: str = 'por') -> str:
    pdfplumber = _import_pdfplumber()
    text_parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = (page.extract_text() or "").strip()
            if _texto_pobre_para_ocr(page_text):
                try:
                    page_img = page.to_image(resolution=250).original
                    ocr_text, _ = _ocr_image(page_img, ocr_lang=ocr_lang)
                    if _preferir_texto_ocr(page_text, ocr_text):
                        page_text = ocr_text
                except Exception:
                    pass
            text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_text_from_docx(path: str) -> str:
    Document = _import_document()
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def extract_text_from_xlsx(path: str) -> str:
    openpyxl = _import_openpyxl()
    wb = openpyxl.load_workbook(path, data_only=True)
    parts = []
    for sheet in wb.worksheets:
        parts.append(f"[Aba: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            row_vals = [str(c) for c in row if c is not None]
            if row_vals:
                parts.append(" | ".join(row_vals))
    return "\n".join(parts)


def extract_text(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == '.pdf':
        return extract_text_from_pdf(path)
    elif ext in ('.png', '.jpg', '.jpeg', '.tif', '.tiff', '.bmp', '.webp'):
        txt, _ = extract_text_from_image(path)
        return txt
    elif ext in ('.docx', '.doc'):
        return extract_text_from_docx(path)
    elif ext in ('.xlsx', '.xls'):
        return extract_text_from_xlsx(path)
    raise ValueError(f"Formato não suportado: {ext}")


def pre_filtrar_texto(text: str) -> str:
    """Remove ruído (parágrafos legais longos sem números, cabeçalhos/rodapés repetidos)
    antes de enviar o texto para a IA — reduz tokens sem descartar itens de verdade.
    Nunca remove uma linha só por não ter preço, para não perder itens não cotados."""
    linhas = [l.strip() for l in text.split("\n")]
    contagem = Counter(l for l in linhas if len(l) > 10)

    linhas_filtradas = []
    repetidas_vistas = Counter()

    for linha in linhas:
        if not linha:
            continue

        tem_digito = any(c.isdigit() for c in linha)
        if len(linha) > BOILERPLATE_MAX_LEN_SEM_DIGITO and not tem_digito:
            continue  # provável parágrafo jurídico/termos, sem nenhum número

        if len(linha) > 10 and contagem[linha] >= 3:
            repetidas_vistas[linha] += 1
            if repetidas_vistas[linha] > 1:
                continue  # cabeçalho/rodapé repetido em várias páginas: mantém só a 1ª ocorrência

        linhas_filtradas.append(linha)

    return "\n".join(linhas_filtradas)


def _parse_numero_br(valor: str):
    if valor is None:
        return None
    texto = str(valor).strip()
    if not texto:
        return None
    texto = re.sub(r"[^\d,.-]", "", texto)
    if not texto:
        return None
    if "," in texto and "." in texto:
        texto = texto.replace(".", "").replace(",", ".")
    elif "," in texto:
        texto = texto.replace(",", ".")
    try:
        return float(texto)
    except ValueError:
        return None


def _clean_unidade(valor: str):
    if valor is None:
        return None
    texto = re.sub(r"\s+", " ", str(valor).strip().upper())
    return texto or None


def _extract_unit_qty_from_segment(segmento: str):
    if not segmento:
        return None, None

    partes = [p.strip() for p in re.split(r"\s*\|\s*", segmento) if p.strip()]
    if len(partes) >= 2:
        for idx, parte in enumerate(partes[:-1]):
            unidade = _clean_unidade(parte)
            if unidade in UNIT_TOKENS:
                quantidade = _parse_numero_br(partes[idx + 1])
                if quantidade is not None:
                    return unidade, quantidade

        for idx, parte in enumerate(partes[1:], start=1):
            unidade = _clean_unidade(parte)
            if unidade in UNIT_TOKENS:
                quantidade = _parse_numero_br(partes[idx - 1])
                if quantidade is not None:
                    return unidade, quantidade

    padroes = [
        (rf"\b({'|'.join(UNIT_TOKENS)})\b\s+(\d+(?:[\.,]\d+)?)\b", "unit_first"),
        (rf"\b(\d+(?:[\.,]\d+)?)\b\s+({'|'.join(UNIT_TOKENS)})\b", "qty_first"),
    ]
    for padrao, ordem in padroes:
        match = re.search(padrao, segmento, flags=re.IGNORECASE)
        if not match:
            continue
        if ordem == "unit_first":
            unidade, quantidade = match.group(1), match.group(2)
        else:
            quantidade, unidade = match.group(1), match.group(2)
        quantidade = _parse_numero_br(quantidade)
        unidade = _clean_unidade(unidade)
        if quantidade is not None and unidade:
            return unidade, quantidade

    return None, None


def _candidate_item_lines(text: str, numero_item, descricao: str):
    linhas = [l.strip() for l in text.splitlines() if l.strip()]
    candidatos = []
    numero = str(numero_item).strip() if numero_item is not None else ""
    descricao_norm = re.sub(r"\s+", " ", (descricao or "").strip().lower())
    palavras_desc = [p for p in re.findall(r"\w+", descricao_norm) if len(p) >= 4][:6]

    for idx, linha in enumerate(linhas):
        linha_norm = linha.lower()
        score = 0
        if numero and re.search(rf"\b{re.escape(numero)}\b", linha):
            score += 5
        score += sum(1 for palavra in palavras_desc if palavra in linha_norm)
        if score > 0:
            candidatos.append((score, idx, linha))

    candidatos.sort(reverse=True)
    resultado = []
    vistos = set()
    for _, idx, linha in candidatos[:5]:
        for near_idx in (idx, idx + 1, idx - 1):
            if 0 <= near_idx < len(linhas) and near_idx not in vistos:
                vistos.add(near_idx)
                resultado.append(linhas[near_idx])
    return resultado


def _enrich_item_fields(result: dict, text: str) -> dict:
    itens = result.get("itens")
    if not isinstance(itens, list):
        return result

    for item in itens:
        if not isinstance(item, dict):
            continue
        precisa_unidade = not item.get("unidade")
        precisa_quantidade = item.get("quantidade") is None
        if not (precisa_unidade or precisa_quantidade):
            continue

        for linha in _candidate_item_lines(text, item.get("numero_item"), item.get("descricao", "")):
            unidade, quantidade = _extract_unit_qty_from_segment(linha)
            if precisa_unidade and unidade:
                item["unidade"] = unidade
                precisa_unidade = False
            if precisa_quantidade and quantidade is not None:
                item["quantidade"] = quantidade
                precisa_quantidade = False
            if not (precisa_unidade or precisa_quantidade):
                break

    return result


# ---------- Interpretação estruturada via LLM (OpenRouter) ----------

EXTRACTION_SYSTEM_PROMPT = """Você é um assistente especializado em extrair itens de orçamentos/cotações de compras.
Dado o texto extraído de um documento de orçamento, retorne APENAS um JSON válido (sem markdown, sem texto adicional) no formato:

{
  "empresa": "nome da empresa/fornecedor, se identificável, senão null",
  "itens": [
    {
      "numero_item": "número do item conforme edital/TR, se houver, senão null",
      "descricao": "descrição do item/produto",
      "unidade": "unidade de medida, se houver, senão null",
      "quantidade": numero ou null,
      "preco_unitario": numero (float, use ponto decimal) ou null,
      "preco_total": numero ou null
    }
  ]
}

Regras:
- Extraia TODOS os itens de orçamento encontrados no texto, inclusive os que aparecem sem preço definido.
- Copie a "descricao" exatamente como está escrita no documento, por extenso — não abrevie, não resuma e não corrija siglas técnicas. Se o mesmo item aparecer mais de uma vez no documento com descrições diferentes (uma mais completa que a outra), use a versão mais completa.
- Quando houver coluna de unidade de fornecimento (UF, UND, UN, PCT, EMB, KG, LT etc.), preencha obrigatoriamente o campo "unidade".
- Quando houver coluna de quantidade (QTD, QUANT, QUANTIDADE), preencha obrigatoriamente o campo "quantidade".
- Em tabelas, preserve a associação da mesma linha do item entre descricao, unidade, quantidade e preco.
- Se preco_unitario nao estiver explicito mas preco_total e quantidade estiverem, calcule preco_unitario = preco_total / quantidade.
- Numeros devem ser float puro, sem simbolo de moeda ou separador de milhar (ex: 1234.56, nunca "R$ 1.234,56").
- Nao invente dados que nao estao no texto.
- Responda em português.
"""


def call_openrouter_extract(text: str, api_key: str, model: str = "anthropic/claude-3.5-sonnet",
                             filename_hint: str = None, max_retries: int = 3,
                             max_chars: int = 50000, pre_filtrar: bool = True) -> dict:
    url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }

    texto_base = pre_filtrar_texto(text) if pre_filtrar else text
    texto_truncado = len(texto_base) > max_chars
    user_content = texto_base[:max_chars]
    if filename_hint:
        user_content = f"Nome do arquivo: {filename_hint}\n\n{user_content}"

    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": EXTRACTION_SYSTEM_PROMPT},
            {"role": "user", "content": user_content},
        ],
        "temperature": 0,
    }

    last_error = "erro desconhecido"
    for attempt in range(max_retries):
        try:
            resp = requests.post(url, headers=headers, json=payload, timeout=120)
            resp.raise_for_status()
            data = resp.json()
            content = data["choices"][0]["message"]["content"]
            content_clean = re.sub(r"^```json\s*|\s*```$", "", content.strip(), flags=re.MULTILINE)
            content_clean = content_clean.strip().strip("`")
            parsed = json.loads(content_clean)
            parsed = _enrich_item_fields(parsed, texto_base)
            if isinstance(parsed.get("itens"), list):
                parsed["itens"] = [normalizar_item(i) for i in parsed["itens"] if isinstance(i, dict)]
            parsed["texto_truncado"] = texto_truncado
            parsed["usage"] = _extract_usage_info(data, model)
            return parsed
        except (requests.RequestException, json.JSONDecodeError, KeyError, IndexError) as exc:
            last_error = str(exc)
            if attempt < max_retries - 1:
                time.sleep(2 ** attempt)

    return {
        "empresa": filename_hint,
        "itens": [],
        "erro": last_error,
        "texto_truncado": texto_truncado,
        "usage": {
            "prompt_tokens": 0,
            "completion_tokens": 0,
            "total_tokens": 0,
            "cost_usd": 0.0,
            "estimated": True,
        },
    }


def _extract_text_pages_pdf(path: str, ocr_lang: str = "por"):
    pdfplumber = _import_pdfplumber()
    pages = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = (page.extract_text() or "").strip()
            veio_ocr = False
            erro_ocr = None
            tentou_ocr = _texto_pobre_para_ocr(page_text)
            if tentou_ocr:
                try:
                    page_img = page.to_image(resolution=250).original
                    ocr_text, erro_ocr = _ocr_image(page_img, ocr_lang=ocr_lang)
                    if _preferir_texto_ocr(page_text, ocr_text):
                        page_text = ocr_text
                        veio_ocr = True
                except Exception as exc:
                    erro_ocr = str(exc)
            pages.append({
                "texto": page_text or "",
                "veio_ocr": veio_ocr,
                "tentou_ocr": tentou_ocr,
                "erro_ocr": erro_ocr,
            })
    return pages


def _inferir_origem_pagina(item: dict, pages: list[dict]) -> str:
    descricao = limpar_quebras_e_caracteres(str(item.get("descricao") or "")).lower()
    numero = str(item.get("numero_item") or "").strip()
    if not pages:
        return "documento"

    for idx, p in enumerate(pages, start=1):
        texto = p.get("texto", "").lower()
        if numero and re.search(rf"\b{re.escape(numero)}\b", texto):
            return f"Pagina {idx}"
        if descricao and descricao[:40] and descricao[:40] in texto:
            return f"Pagina {idx}"

    return "Pagina 1"


def _anotar_fonte_origem(itens: list[dict], fonte: str, pages: list[dict] | None = None) -> list[dict]:
    saida = []
    for item in itens or []:
        i2 = dict(item)
        i2["fonte_extracao"] = fonte
        if not i2.get("origem"):
            i2["origem"] = _inferir_origem_pagina(i2, pages or [])
        saida.append(normalizar_item(i2))
    return saida


def calcular_similaridade_item(item_parser: dict, item_ia: dict) -> float:
    num_parser = str(item_parser.get("numero_item") or "").strip()
    num_ia = str(item_ia.get("numero_item") or "").strip()
    if num_parser and num_ia and num_parser == num_ia:
        return 100.0

    desc_parser = limpar_quebras_e_caracteres(str(item_parser.get("descricao") or "")).lower()
    desc_ia = limpar_quebras_e_caracteres(str(item_ia.get("descricao") or "")).lower()
    sim_desc = token_set_ratio(desc_parser, desc_ia)

    unidade_parser = str(item_parser.get("unidade") or "").upper()
    unidade_ia = str(item_ia.get("unidade") or "").upper()
    sim_unidade = 100.0 if unidade_parser and unidade_parser == unidade_ia else 0.0

    qtd_p = item_parser.get("quantidade")
    qtd_i = item_ia.get("quantidade")
    sim_qtd = 0.0
    if qtd_p is not None and qtd_i is not None:
        if qtd_p == qtd_i:
            sim_qtd = 100.0
        elif max(abs(qtd_p), abs(qtd_i)) > 0:
            diff = abs(qtd_p - qtd_i) / max(abs(qtd_p), abs(qtd_i))
            if diff < 0.05:
                sim_qtd = 100.0

    return (0.60 * sim_desc) + (0.25 * sim_unidade) + (0.15 * sim_qtd)


def comparar_extracoes(itens_parser: list[dict], itens_ia: list[dict]):
    def _mesclar_item(parser_item: dict, ia_item: dict, score: float, numerico_bate: bool) -> dict:
        item_final = dict(parser_item)
        item_final["fonte_extracao"] = "dupla_checagem"
        item_final["confianca"] = "alta" if score >= 70 and numerico_bate else "baixa"

        if ia_item.get("descricao") and len(str(ia_item.get("descricao"))) > len(str(item_final.get("descricao") or "")):
            item_final["descricao"] = ia_item.get("descricao")

        for campo in ("unidade", "quantidade", "preco_total", "origem"):
            if item_final.get(campo) in (None, "") and ia_item.get(campo) not in (None, ""):
                item_final[campo] = ia_item.get(campo)

        if item_final.get("preco_unitario") is None and ia_item.get("preco_unitario") is not None:
            item_final["preco_unitario"] = ia_item.get("preco_unitario")

        return normalizar_item(item_final)

    pares = []
    for i, item_p in enumerate(itens_parser):
        for j, item_i in enumerate(itens_ia):
            score = calcular_similaridade_item(item_p, item_i)
            pares.append((score, i, j))

    pares.sort(reverse=True, key=lambda x: x[0])
    usados_parser = set()
    usados_ia = set()
    casamentos = []
    for score, i, j in pares:
        if i in usados_parser or j in usados_ia:
            continue
        usados_parser.add(i)
        usados_ia.add(j)
        casamentos.append((score, itens_parser[i], itens_ia[j]))

    itens_finais = []
    conflitos = []
    for score, parser_item, ia_item in casamentos:
        p_unit = parser_item.get("preco_unitario")
        i_unit = ia_item.get("preco_unitario")
        numerico_bate = False
        if p_unit is not None and i_unit is not None and max(abs(p_unit), abs(i_unit), 1e-9) > 0:
            numerico_bate = abs(p_unit - i_unit) / max(abs(p_unit), abs(i_unit)) < 0.01

        itens_finais.append(_mesclar_item(parser_item, ia_item, score, numerico_bate))

        if score < 70 or not numerico_bate:
            conflitos.append({
                "tipo": "divergência parser vs. IA",
                "numero_item": parser_item.get("numero_item") or ia_item.get("numero_item"),
                "descricao_nova": parser_item.get("descricao") or "",
                "casou_com": ia_item.get("descricao") or "",
                "score": round(score, 2),
            })

    for idx, parser_item in enumerate(itens_parser):
        if idx not in usados_parser:
            item_final = dict(parser_item)
            item_final["fonte_extracao"] = "dupla_checagem"
            item_final["confianca"] = "baixa"
            itens_finais.append(normalizar_item(item_final))
            conflitos.append({
                "tipo": "divergência parser vs. IA",
                "numero_item": parser_item.get("numero_item"),
                "descricao_nova": parser_item.get("descricao") or "",
                "casou_com": "sem par na IA",
                "score": 0,
            })

    for idx, ia_item in enumerate(itens_ia):
        if idx not in usados_ia:
            item_final = dict(ia_item)
            item_final["fonte_extracao"] = "dupla_checagem"
            item_final["confianca"] = "média"
            itens_finais.append(normalizar_item(item_final))
            conflitos.append({
                "tipo": "divergência parser vs. IA",
                "numero_item": ia_item.get("numero_item"),
                "descricao_nova": ia_item.get("descricao") or "",
                "casou_com": "sem par no parser",
                "score": 0,
            })

    return itens_finais, conflitos


def extrair_orcamento_em_camadas(path: str, api_key: str, model: str,
                                 pre_filtrar: bool = True,
                                 limiar_alto: int = 85,
                                 limiar_baixo: int = 40) -> dict:
    """Pipeline em camadas: estrutural quando possivel, IA por excecao."""
    tipo = detectar_tipo_e_rotear(path)
    nome_arquivo = os.path.basename(path)
    review = []
    confianca = None
    texto_base = ""
    texto_truncado = False
    fonte_global = "ia"
    debug_events = []
    debug_highlights = []

    debug_events.append(f"Arquivo recebido: {nome_arquivo}")
    debug_events.append(f"Percepcao inicial: tipo detectado = {tipo}")

    if tipo == "xlsx":
        debug_events.append("Biblioteca acionada: openpyxl (extracao estrutural de planilha)")
    elif tipo == "docx":
        debug_events.append("Biblioteca acionada: python-docx (extracao estrutural de documento Word)")
    elif tipo in ("pdf_texto", "pdf_escaneado"):
        debug_events.append("Biblioteca acionada: pdfplumber (leitura de paginas PDF)")
        debug_events.append("Biblioteca potencial: requests (chamada IA quando necessario)")
    elif tipo == "imagem":
        debug_events.append("Biblioteca acionada: Pillow + pytesseract (OCR de imagem)")
        debug_events.append("Biblioteca potencial: requests (chamada IA quando necessario)")

    if tipo == "xlsx":
        itens = extrair_xlsx_estruturado(path)
        if itens:
            fonte_global = "estrutural"
            texto_base = extract_text(path)
            empresa_det = extrair_razao_social(texto_base)
            debug_events.append(f"Achado: {len(itens)} item(ns) extraido(s) por modo estrutural")
            debug_highlights.append("RELEVANTE: processamento 100% estrutural (sem consumo de IA)")
            return {
                "empresa": empresa_det or os.path.splitext(nome_arquivo)[0],
                "itens": itens,
                "review": review,
                "confianca_estrutural": 100,
                "fonte_processamento": fonte_global,
                "texto_truncado": False,
                "debug_events": debug_events,
                "debug_highlights": debug_highlights,
                "usage": {
                    "prompt_tokens": 0,
                    "completion_tokens": 0,
                    "total_tokens": 0,
                    "cost_usd": 0.0,
                    "estimated": False,
                },
            }

    if tipo == "docx":
        itens = extrair_docx_estruturado(path)
        if itens:
            fonte_global = "estrutural"
            texto_base = extract_text(path)
            empresa_det = extrair_razao_social(texto_base)
            debug_events.append(f"Achado: {len(itens)} item(ns) extraido(s) por modo estrutural")
            debug_highlights.append("RELEVANTE: processamento 100% estrutural (sem consumo de IA)")
            return {
                "empresa": empresa_det or os.path.splitext(nome_arquivo)[0],
                "itens": itens,
                "review": review,
                "confianca_estrutural": 100,
                "fonte_processamento": fonte_global,
                "texto_truncado": False,
                "debug_events": debug_events,
                "debug_highlights": debug_highlights,
                "usage": {
                    "prompt_tokens": 0,
                    "completion_tokens": 0,
                    "total_tokens": 0,
                    "cost_usd": 0.0,
                    "estimated": False,
                },
            }

    if tipo in ("pdf_texto", "pdf_escaneado"):
        pages = _extract_text_pages_pdf(path)
        texto_base = "\n".join(p["texto"] for p in pages)
        debug_events.append(f"Achado: {len(pages)} pagina(s) lida(s) no PDF")
        debug_events.append(f"Achado: {len(texto_base)} caractere(s) extraido(s) do documento")
        paginas_ocr = sum(1 for p in pages if p.get("veio_ocr"))
        paginas_tentou_ocr = sum(1 for p in pages if p.get("tentou_ocr"))
        if paginas_tentou_ocr:
            debug_events.append(f"Achado: OCR tentado em {paginas_tentou_ocr} pagina(s)")
        if paginas_ocr:
            debug_events.append(f"Achado: OCR aplicado em {paginas_ocr} pagina(s)")

        if not texto_base.strip():
            erros_ocr = [p.get("erro_ocr") for p in pages if p.get("erro_ocr")]
            mensagem_erro = "Texto vazio apos leitura do PDF; nao foi possivel extrair conteudo util."
            if erros_ocr:
                mensagem_erro = (
                    "Texto vazio apos leitura do PDF e OCR indisponivel/falhou. "
                    f"Detalhe: {erros_ocr[0]}"
                )
                debug_highlights.append("RELEVANTE: OCR nao disponivel no ambiente")
            return {
                "empresa": os.path.splitext(nome_arquivo)[0],
                "itens": [],
                "erro": mensagem_erro,
                "review": review,
                "confianca_estrutural": 0,
                "fonte_processamento": "ia",
                "texto_truncado": False,
                "cnpj": None,
                "debug_events": debug_events,
                "debug_highlights": debug_highlights,
                "usage": {
                    "prompt_tokens": 0,
                    "completion_tokens": 0,
                    "total_tokens": 0,
                    "cost_usd": 0.0,
                    "estimated": False,
                },
            }
        resultado_estrutural = tentar_extracao_estrutural_pdf(path)
        confianca = calcular_confianca_estrutural(resultado_estrutural)
        if tipo == "pdf_escaneado":
            confianca = max(0, confianca - 10)
            debug_events.append("Ajuste aplicado: PDF escaneado detectado, penalidade de confianca estrutural")

        debug_events.append(f"Percepcao do sistema: confianca estrutural = {confianca}")

        if confianca >= limiar_alto and resultado_estrutural.get("itens"):
            itens_finais = _anotar_fonte_origem(resultado_estrutural["itens"], "estrutural", pages)
            fonte_global = "estrutural"
            empresa_det = extrair_razao_social(texto_base)
            debug_events.append(f"Achado: {len(itens_finais)} item(ns) por parser estrutural")
            debug_highlights.append(
                f"RELEVANTE: confianca estrutural alta ({confianca}) -> IA nao foi necessaria"
            )
            return {
                "empresa": empresa_det or os.path.splitext(nome_arquivo)[0],
                "itens": itens_finais,
                "review": review,
                "confianca_estrutural": confianca,
                "fonte_processamento": fonte_global,
                "texto_truncado": False,
                "debug_events": debug_events,
                "debug_highlights": debug_highlights,
                "usage": {
                    "prompt_tokens": 0,
                    "completion_tokens": 0,
                    "total_tokens": 0,
                    "cost_usd": 0.0,
                    "estimated": False,
                },
            }

        llm = call_openrouter_extract(
            texto_base,
            api_key,
            model=model,
            filename_hint=nome_arquivo,
            pre_filtrar=pre_filtrar,
        )
        debug_events.append(f"Biblioteca acionada: requests (modelo IA = {model})")
        texto_truncado = bool(llm.get("texto_truncado"))
        if texto_truncado:
            debug_highlights.append("RELEVANTE: texto truncado para caber no limite enviado a IA")
        itens_ia = _anotar_fonte_origem(llm.get("itens", []), "ia", pages)
        debug_events.append(f"Achado IA: {len(itens_ia)} item(ns) retornado(s)")

        if confianca is not None and confianca < limiar_baixo:
            fonte_global = "ia"
            itens_finais = itens_ia
            debug_highlights.append(
                f"RELEVANTE: confianca estrutural baixa ({confianca}) -> decisao final por IA"
            )
        elif resultado_estrutural.get("itens"):
            fonte_global = "dupla_checagem"
            itens_finais, conflitos = comparar_extracoes(resultado_estrutural["itens"], itens_ia)
            review.extend(conflitos)
            debug_events.append(
                f"Percepcao: dupla checagem executada (parser vs IA), conflitos={len(conflitos)}"
            )
            if conflitos:
                debug_highlights.append(
                    f"RELEVANTE: divergencias detectadas entre parser e IA ({len(conflitos)})"
                )
        else:
            fonte_global = "ia"
            itens_finais = itens_ia
            debug_events.append("Percepcao: sem base estrutural valida, decisao final por IA")

        empresa_det = extrair_razao_social(texto_base)
        empresa_llm = llm.get("empresa")
        if empresa_llm:
            debug_events.append(f"Achado: empresa identificada pela IA = {empresa_llm}")
        elif empresa_det:
            debug_events.append(f"Achado: empresa inferida por heuristica = {empresa_det}")

        if extrair_cnpj(texto_base):
            debug_events.append("Achado: CNPJ identificado no documento")

        return {
            "empresa": empresa_llm or empresa_det or os.path.splitext(nome_arquivo)[0],
            "itens": [normalizar_item(i) for i in itens_finais],
            "review": review,
            "confianca_estrutural": confianca,
            "fonte_processamento": fonte_global,
            "texto_truncado": texto_truncado,
            "cnpj": extrair_cnpj(texto_base),
            "debug_events": debug_events,
            "debug_highlights": debug_highlights,
            "usage": llm.get("usage", {
                "prompt_tokens": 0,
                "completion_tokens": 0,
                "total_tokens": 0,
                "cost_usd": 0.0,
                "estimated": True,
            }),
        }

    if tipo == "imagem":
        texto_base, erro_ocr = extract_text_from_image(path)
        debug_events.append(f"Achado: {len(texto_base)} caractere(s) extraido(s) da imagem")
        if erro_ocr and not texto_base.strip():
            debug_highlights.append("RELEVANTE: OCR nao disponivel no ambiente")
            return {
                "empresa": os.path.splitext(nome_arquivo)[0],
                "itens": [],
                "erro": f"Falha no OCR da imagem: {erro_ocr}",
                "review": review,
                "confianca_estrutural": 0,
                "fonte_processamento": "ia",
                "texto_truncado": False,
                "cnpj": None,
                "debug_events": debug_events,
                "debug_highlights": debug_highlights,
                "usage": {
                    "prompt_tokens": 0,
                    "completion_tokens": 0,
                    "total_tokens": 0,
                    "cost_usd": 0.0,
                    "estimated": False,
                },
            }

        if not texto_base.strip():
            return {
                "empresa": os.path.splitext(nome_arquivo)[0],
                "itens": [],
                "erro": "A imagem nao possui texto legivel para extracao.",
                "review": review,
                "confianca_estrutural": 0,
                "fonte_processamento": "ia",
                "texto_truncado": False,
                "cnpj": None,
                "debug_events": debug_events,
                "debug_highlights": debug_highlights,
                "usage": {
                    "prompt_tokens": 0,
                    "completion_tokens": 0,
                    "total_tokens": 0,
                    "cost_usd": 0.0,
                    "estimated": False,
                },
            }

        llm = call_openrouter_extract(
            texto_base,
            api_key,
            model=model,
            filename_hint=nome_arquivo,
            pre_filtrar=pre_filtrar,
        )
        debug_events.append(f"Biblioteca acionada: requests (modelo IA = {model})")
        debug_events.append(f"Achado IA: {len(llm.get('itens', []))} item(ns) retornado(s)")
        itens = _anotar_fonte_origem(llm.get("itens", []), "ia")
        empresa_det = extrair_razao_social(texto_base)
        debug_highlights.append("RELEVANTE: imagem processada por OCR + IA")

        return {
            "empresa": llm.get("empresa") or empresa_det or os.path.splitext(nome_arquivo)[0],
            "itens": itens,
            "review": review,
            "confianca_estrutural": confianca,
            "fonte_processamento": "ia",
            "texto_truncado": bool(llm.get("texto_truncado")),
            "cnpj": extrair_cnpj(texto_base),
            "debug_events": debug_events,
            "debug_highlights": debug_highlights,
            "usage": llm.get("usage", {
                "prompt_tokens": 0,
                "completion_tokens": 0,
                "total_tokens": 0,
                "cost_usd": 0.0,
                "estimated": True,
            }),
        }

    texto_base = extract_text(path)
    llm = call_openrouter_extract(
        texto_base,
        api_key,
        model=model,
        filename_hint=nome_arquivo,
        pre_filtrar=pre_filtrar,
    )
    debug_events.append(f"Biblioteca acionada: requests (modelo IA = {model})")
    debug_events.append(f"Achado IA: {len(llm.get('itens', []))} item(ns) retornado(s)")
    itens = _anotar_fonte_origem(llm.get("itens", []), "ia")
    empresa_det = extrair_razao_social(texto_base)
    debug_highlights.append("RELEVANTE: arquivo roteado diretamente para IA")

    return {
        "empresa": llm.get("empresa") or empresa_det or os.path.splitext(nome_arquivo)[0],
        "itens": itens,
        "review": review,
        "confianca_estrutural": confianca,
        "fonte_processamento": fonte_global,
        "texto_truncado": bool(llm.get("texto_truncado")),
        "cnpj": extrair_cnpj(texto_base),
        "debug_events": debug_events,
        "debug_highlights": debug_highlights,
        "usage": llm.get("usage", {
            "prompt_tokens": 0,
            "completion_tokens": 0,
            "total_tokens": 0,
            "cost_usd": 0.0,
            "estimated": True,
        }),
    }
