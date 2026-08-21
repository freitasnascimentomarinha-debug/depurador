"""
Extração de texto de PDF (com OCR para escaneados), Word e Excel,
e interpretação estruturada via LLM (OpenRouter).
"""
from __future__ import annotations

import hashlib
import io
import json
import os
import re
import shutil
import threading
import time
import csv
import unicodedata
from concurrent.futures import ThreadPoolExecutor
from collections import Counter

import requests

from http_client import SESSAO

from confidence import calcular_confianca_estrutural
from normalize_utils import (
    extrair_cnpj,
    extrair_telefone,
    extrair_razao_social,
    limpar_quebras_e_caracteres,
    normalizar_item,
)
from structured_extract import (
    _import_openpyxl,
    detectar_tipo_e_rotear,
    extrair_docx_estruturado,
    extrair_xlsx_estruturado,
    tentar_extracao_estrutural_pdf,
)
from text_similarity import token_set_ratio

BOILERPLATE_MAX_LEN_SEM_DIGITO = 150
UNIT_TOKENS = (
    "UN", "UND", "UNID", "UNIDADE", "PCT", "PC", "PÇ", "EMB", "KG", "G", "MG",
    "L", "LT", "ML", "CX", "FR", "FD", "RL", "M", "M2", "M3", "PAR", "CJ",
)

# Precos aproximados por 1M tokens (USD) — usados apenas como fallback de estimativa.
# O custo real vem da propria API via payload {"usage": {"include": true}}.
MODEL_PRICING_PER_MILLION = {
    "google/gemini-2.5-flash": {"input": 0.30, "output": 2.50},
    "openai/gpt-5-mini": {"input": 0.25, "output": 2.00},
    "deepseek/deepseek-chat-v3.2": {"input": 0.25, "output": 0.40},
    # legados (mantidos para compatibilidade com cache antigo)
    "openai/gpt-4o-mini": {"input": 0.15, "output": 0.60},
    "anthropic/claude-3-haiku": {"input": 0.25, "output": 1.25},
}

# Modelo padrao de extracao: rapido, com suporte a structured outputs.
DEFAULT_EXTRACTION_MODEL = "qwen/qwen3.7-flash:nitro"

# Escalonamento: quando o modelo barato falha ou extrai zero itens, uma unica
# retentativa com modelo forte (pago so nos ~5% de documentos dificeis).
ESCALATION_MODEL = "openai/gpt-5.6-luna"

_OCR_MIN_ALNUM = 25
_OCR_MIN_TOKENS = 6
_PAGINAS_CACHE: dict[str, list[dict]] = {}
_PAGINAS_CACHE_LOCK = threading.Lock()
_OCR_SEMAFORO = threading.BoundedSemaphore(max(1, (os.cpu_count() or 4)))


def _chave_arquivo(path: str) -> str:
    st_ = os.stat(path)
    return f"{os.path.abspath(path)}|{st_.st_size}|{st_.st_mtime_ns}"


def _sha256_arquivo(path: str, chunk: int = 1 << 20) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as fh:
        for bloco in iter(lambda: fh.read(chunk), b""):
            h.update(bloco)
    return h.hexdigest()


def _to_int(value) -> int:
    try:
        return int(value)
    except (TypeError, ValueError):
        return 0


def _extract_usage_info(response_json: dict, model: str) -> dict:
    usage = response_json.get("usage") or {}
    prompt_tokens = _to_int(usage.get("prompt_tokens") or usage.get("input_tokens"))
    completion_tokens = _to_int(usage.get("completion_tokens") or usage.get("output_tokens"))
    total_tokens = _to_int(usage.get("total_tokens"))
    if total_tokens <= 0:
        total_tokens = prompt_tokens + completion_tokens

    cost_usd = usage.get("cost")
    if cost_usd is None:
        cost_usd = usage.get("total_cost")

    estimated = False
    if cost_usd is None:
        pricing = MODEL_PRICING_PER_MILLION.get(model)
        if pricing:
            cost_usd = (
                (prompt_tokens / 1_000_000.0) * pricing["input"]
                + (completion_tokens / 1_000_000.0) * pricing["output"]
            )
            estimated = True
        else:
            cost_usd = 0.0
            estimated = True

    try:
        cost_usd = float(cost_usd)
    except (TypeError, ValueError):
        cost_usd = 0.0

    return {
        "prompt_tokens": prompt_tokens,
        "completion_tokens": completion_tokens,
        "total_tokens": total_tokens,
        "cost_usd": cost_usd,
        "estimated": estimated,
    }


def _import_pytesseract():
    import pytesseract
    return pytesseract


def _import_image():
    from PIL import Image
    return Image


def _import_document():
    from docx import Document
    return Document


def _import_openpyxl():
    import openpyxl
    return openpyxl


def _import_pdfplumber():
    import pdfplumber
    return pdfplumber


def _ocr_lang_candidates(ocr_lang: str) -> list[str]:
    base = (ocr_lang or "por").strip()
    candidates = [base]
    if "+" not in base:
        candidates.append(f"{base}+eng")
    if "eng" not in candidates:
        candidates.append("eng")
    vistos = set()
    unicos = []
    for cand in candidates:
        if cand and cand not in vistos:
            vistos.add(cand)
            unicos.append(cand)
    return unicos


def ocr_runtime_status() -> tuple[bool, str | None]:
    """Valida se o OCR esta operacional no ambiente atual."""
    if shutil.which("tesseract") is None:
        return False, "Binario 'tesseract' nao encontrado no PATH."

    try:
        pytesseract = _import_pytesseract()
        _ = pytesseract.get_tesseract_version()
    except Exception as exc:
        return False, f"Falha ao inicializar pytesseract/tesseract: {exc}"

    return True, None


def _texto_pobre_para_ocr(texto: str) -> bool:
    """Indica se o texto extraido da pagina esta fraco e deve tentar OCR."""
    base = re.sub(r"\s+", " ", str(texto or "")).strip()
    if not base:
        return True

    alnum = sum(1 for ch in base if ch.isalnum())
    tokens = re.findall(r"\w+", base)
    if alnum < _OCR_MIN_ALNUM:
        return True
    if len(tokens) < _OCR_MIN_TOKENS:
        return True
    return False


def _preferir_texto_ocr(texto_pdf: str, texto_ocr: str) -> bool:
    """Troca para OCR quando ele traz ganho claro de conteudo util."""
    base_pdf = re.sub(r"\s+", " ", str(texto_pdf or "")).strip()
    base_ocr = re.sub(r"\s+", " ", str(texto_ocr or "")).strip()
    if not base_ocr:
        return False
    if not base_pdf:
        return True
    return len(base_ocr) >= int(len(base_pdf) * 1.2)


def _ocr_image(image_obj, ocr_lang: str = "por") -> tuple[str, str | None]:
    with _OCR_SEMAFORO:
        ok_ocr, erro_ocr = ocr_runtime_status()
        if not ok_ocr:
            return "", erro_ocr

        pytesseract = _import_pytesseract()
        last_error = None
        for lang in _ocr_lang_candidates(ocr_lang):
            try:
                texto = pytesseract.image_to_string(image_obj, lang=lang)
                if texto and texto.strip():
                    return texto.strip(), None
            except Exception as exc:
                last_error = str(exc)
        return "", last_error


def extract_text_from_image(path: str, ocr_lang: str = 'por') -> tuple[str, str | None]:
    Image = _import_image()
    with Image.open(path) as img:
        return _ocr_image(img, ocr_lang=ocr_lang)


# ---------- Extração de texto bruto ----------

def extract_text_from_pdf(path: str, ocr_lang: str = 'por') -> str:
    pdfplumber = _import_pdfplumber()
    text_parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = (page.extract_text() or "").strip()
            if _texto_pobre_para_ocr(page_text):
                try:
                    page_img = page.to_image(resolution=250).original
                    ocr_text, _ = _ocr_image(page_img, ocr_lang=ocr_lang)
                    if _preferir_texto_ocr(page_text, ocr_text):
                        page_text = ocr_text
                except Exception:
                    pass
            text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_text_from_docx(path: str) -> str:
    Document = _import_document()
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def extract_text_from_xlsx(path: str) -> str:
    openpyxl = _import_openpyxl()
    wb = openpyxl.load_workbook(path, data_only=True)
    parts = []
    for sheet in wb.worksheets:
        parts.append(f"[Aba: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            row_vals = [str(c) for c in row if c is not None]
            if row_vals:
                parts.append(" | ".join(row_vals))
    return "\n".join(parts)


def extract_text_from_csv(path: str) -> str:
    parts = []
    with open(path, "r", encoding="utf-8", errors="ignore", newline="") as fh:
        sample = fh.read(4096)
        fh.seek(0)
        try:
            dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
            reader = csv.reader(fh, dialect)
        except Exception:
            reader = csv.reader(fh)
        for row in reader:
            vals = [str(c).strip() for c in row if str(c).strip()]
            if vals:
                parts.append(" | ".join(vals))
    return "\n".join(parts)


def extract_text(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == '.pdf':
        return extract_text_from_pdf(path)
    elif ext in ('.png', '.jpg', '.jpeg', '.tif', '.tiff', '.bmp', '.webp'):
        txt, _ = extract_text_from_image(path)
        return txt
    elif ext in ('.docx', '.doc'):
        return extract_text_from_docx(path)
    elif ext in ('.xlsx', '.xls'):
        return extract_text_from_xlsx(path)
    elif ext == '.csv':
        return extract_text_from_csv(path)
    raise ValueError(f"Formato não suportado: {ext}")


def pre_filtrar_texto(text: str) -> str:
    """Remove ruído (parágrafos legais longos sem números, cabeçalhos/rodapés repetidos)
    antes de enviar o texto para a IA — reduz tokens sem descartar itens de verdade.
    Nunca remove uma linha só por não ter preço, para não perder itens não cotados."""
    linhas = [l.strip() for l in text.split("\n")]
    contagem = Counter(l for l in linhas if len(l) > 10)

    linhas_filtradas = []
    repetidas_vistas = Counter()

    for linha in linhas:
        if not linha:
            continue

        tem_digito = any(c.isdigit() for c in linha)
        if len(linha) > BOILERPLATE_MAX_LEN_SEM_DIGITO and not tem_digito:
            continue  # provável parágrafo jurídico/termos, sem nenhum número

        if len(linha) > 10 and contagem[linha] >= 3:
            repetidas_vistas[linha] += 1
            if repetidas_vistas[linha] > 1:
                continue  # cabeçalho/rodapé repetido em várias páginas: mantém só a 1ª ocorrência

        linhas_filtradas.append(linha)

    return "\n".join(linhas_filtradas)


def _parse_numero_br(valor: str):
    if valor is None:
        return None
    texto = str(valor).strip()
    if not texto:
        return None
    texto = re.sub(r"[^\d,.-]", "", texto)
    if not texto:
        return None
    if "," in texto and "." in texto:
        texto = texto.replace(".", "").replace(",", ".")
    elif "," in texto:
        texto = texto.replace(",", ".")
    try:
        return float(texto)
    except ValueError:
        return None


def _clean_unidade(valor: str):
    if valor is None:
        return None
    texto = re.sub(r"\s+", " ", str(valor).strip().upper())
    return texto or None


def _extract_unit_qty_from_segment(segmento: str):
    if not segmento:
        return None, None

    partes = [p.strip() for p in re.split(r"\s*\|\s*", segmento) if p.strip()]
    if len(partes) >= 2:
        for idx, parte in enumerate(partes[:-1]):
            unidade = _clean_unidade(parte)
            if unidade in UNIT_TOKENS:
                quantidade = _parse_numero_br(partes[idx + 1])
                if quantidade is not None:
                    return unidade, quantidade

        for idx, parte in enumerate(partes[1:], start=1):
            unidade = _clean_unidade(parte)
            if unidade in UNIT_TOKENS:
                quantidade = _parse_numero_br(partes[idx - 1])
                if quantidade is not None:
                    return unidade, quantidade

    padroes = [
        (rf"\b({'|'.join(UNIT_TOKENS)})\b\s+(\d+(?:[\.,]\d+)?)\b", "unit_first"),
        (rf"\b(\d+(?:[\.,]\d+)?)\b\s+({'|'.join(UNIT_TOKENS)})\b", "qty_first"),
    ]
    for padrao, ordem in padroes:
        match = re.search(padrao, segmento, flags=re.IGNORECASE)
        if not match:
            continue
        if ordem == "unit_first":
            unidade, quantidade = match.group(1), match.group(2)
        else:
            quantidade, unidade = match.group(1), match.group(2)
        quantidade = _parse_numero_br(quantidade)
        unidade = _clean_unidade(unidade)
        if quantidade is not None and unidade:
            return unidade, quantidade

    return None, None


def _candidate_item_lines(text: str, numero_item, descricao: str):
    linhas = [l.strip() for l in text.splitlines() if l.strip()]
    candidatos = []
    numero = str(numero_item).strip() if numero_item is not None else ""
    descricao_norm = re.sub(r"\s+", " ", (descricao or "").strip().lower())
    palavras_desc = [p for p in re.findall(r"\w+", descricao_norm) if len(p) >= 4][:6]

    for idx, linha in enumerate(linhas):
        linha_norm = linha.lower()
        score = 0
        if numero and re.search(rf"\b{re.escape(numero)}\b", linha):
            score += 5
        score += sum(1 for palavra in palavras_desc if palavra in linha_norm)
        if score > 0:
            candidatos.append((score, idx, linha))

    candidatos.sort(reverse=True)
    resultado = []
    vistos = set()
    for _, idx, linha in candidatos[:5]:
        for near_idx in (idx, idx + 1, idx - 1):
            if 0 <= near_idx < len(linhas) and near_idx not in vistos:
                vistos.add(near_idx)
                resultado.append(linhas[near_idx])
    return resultado


def _enrich_item_fields(result: dict, text: str) -> dict:
    if not isinstance(result, dict):
        return result
    itens = result.get("itens")
    if not isinstance(itens, list):
        return result

    for item in itens:
        if not isinstance(item, dict):
            continue
        precisa_unidade = not item.get("unidade")
        precisa_quantidade = item.get("quantidade") is None
        if not (precisa_unidade or precisa_quantidade):
            continue

        for linha in _candidate_item_lines(text, item.get("numero_item"), item.get("descricao", "")):
            unidade, quantidade = _extract_unit_qty_from_segment(linha)
            if precisa_unidade and unidade:
                item["unidade"] = unidade
                precisa_unidade = False
            if precisa_quantidade and quantidade is not None:
                item["quantidade"] = quantidade
                precisa_quantidade = False
            if not (precisa_unidade or precisa_quantidade):
                break

    return result


# ---------- Interpretação estruturada via LLM (OpenRouter) ----------

EXTRACTION_SYSTEM_PROMPT = """Você é um assistente especializado em extrair itens de orçamentos/cotações de compras.
Dado o texto extraído de um documento de orçamento, retorne APENAS um JSON válido (sem markdown, sem texto adicional) no formato:

{
  "empresa": "nome da empresa/fornecedor, se identificável, senão null",
  "itens": [
    {
      "numero_item": "número do item conforme edital/TR, se houver, senão null",
      "codigo": "código único do item (PI, NSN, Part Number, Nº de Estoque, Código do Item), se houver, senão null",
      "descricao": "descrição do item/produto",
      "unidade": "unidade de medida, se houver, senão null",
      "quantidade": numero ou null,
      "preco_unitario": numero (float, use ponto decimal) ou null,
      "preco_total": numero ou null
    }
  ]
}

Regras:
- "empresa" deve conter somente um nome curto para identificar o fornecedor no cabeçalho do mapa comparativo: prefira a marca ou a primeira palavra distintiva (ex.: "AKITA", "NEXBOLT", "ORBITAL"). Nunca inclua CNPJ, e-mail, endereço, item, UF, código, tipo societário (LTDA, ME, EPP etc.) ou a razão social completa.
- Extraia TODOS os itens de orçamento encontrados no texto, inclusive os que aparecem sem preço definido.
- "codigo" é o identificador único do MATERIAL (colunas tipo PI, NSN, P/N, Part Number, Nº Estoque, Cód. Item, Referência) — não confundir com "numero_item", que é a posição sequencial do item no edital (1, 2, 3...). Copie o código exatamente como está, com traços e pontos.
- Copie a "descricao" exatamente como está escrita no documento, por extenso — não abrevie, não resuma e não corrija siglas técnicas. Se o mesmo item aparecer mais de uma vez no documento com descrições diferentes (uma mais completa que a outra), use a versão mais completa.
- Quando houver coluna de unidade de fornecimento (UF, UND, UN, PCT, EMB, KG, LT etc.), preencha obrigatoriamente o campo "unidade".
- Quando houver coluna de quantidade (QTD, QUANT, QUANTIDADE), preencha obrigatoriamente o campo "quantidade".
- Em tabelas, preserve a associação da mesma linha do item entre descricao, unidade, quantidade e preco.
- Se preco_unitario nao estiver explicito mas preco_total e quantidade estiverem, calcule preco_unitario = preco_total / quantidade.
- Numeros devem ser float puro, sem simbolo de moeda ou separador de milhar (ex: 1234.56, nunca "R$ 1.234,56").
- Nao invente dados que nao estao no texto.
- Responda em português.

- Se for fornecida uma "LISTA DE REFERÊNCIA" (itens da solicitação de orçamento/edital, com numero_item e codigo oficiais), use-a para identificar a qual item de referência cada item do orçamento corresponde, mesmo quando o fornecedor escreve a descrição com palavras diferentes, abreviada ou incompleta. Quando encontrar correspondência, preencha "numero_item" e "codigo" EXATAMENTE como estão na lista de referência (não invente numeração própria do fornecedor se ela divergir da lista de referência). Se o item do orçamento não corresponder a nenhum item da lista de referência, ainda assim extraia-o normalmente, com numero_item/codigo null se não houver correspondência clara.
"""


# JSON Schema para structured outputs (OpenRouter response_format).
# Garante JSON valido e tipado direto do modelo, sem parsing fragil de markdown.
EXTRACTION_JSON_SCHEMA = {
    "name": "extracao_orcamento",
    "strict": True,
    "schema": {
        "type": "object",
        "properties": {
            "empresa": {"type": ["string", "null"]},
            "itens": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "numero_item": {"type": ["string", "null"]},
                        "codigo": {"type": ["string", "null"]},
                        "descricao": {"type": "string"},
                        "unidade": {"type": ["string", "null"]},
                        "quantidade": {"type": ["number", "null"]},
                        "preco_unitario": {"type": ["number", "null"]},
                        "preco_total": {"type": ["number", "null"]},
                    },
                    "required": ["numero_item", "codigo", "descricao", "unidade",
                                 "quantidade", "preco_unitario", "preco_total"],
                    "additionalProperties": False,
                },
            },
        },
        "required": ["empresa", "itens"],
        "additionalProperties": False,
    },
}


def _formatar_lista_referencia(lista_referencia: list[dict] | None) -> str:
    """Formata a lista de itens da solicitação de orçamento/edital (numero_item,
    codigo, descricao) como bloco de contexto para a IA usar no casamento.
    Nunca deixa um item malformado (não-dicionário) derrubar a extração
    inteira: ignora silenciosamente entradas que não sejam dict."""
    if not lista_referencia:
        return ""
    linhas = ["LISTA DE REFERÊNCIA (itens oficiais da solicitação de orçamento/edital):"]
    for it in lista_referencia:
        if not isinstance(it, dict):
            continue
        numero = it.get("numero_item")
        codigo = it.get("codigo")
        descricao = it.get("descricao") or ""
        partes = []
        if numero is not None:
            partes.append(f"nº {numero}")
        if codigo:
            partes.append(f"código {codigo}")
        partes.append(descricao)
        linhas.append("- " + " | ".join(str(p) for p in partes if p))
    if len(linhas) <= 1:
        return ""
    return "\n".join(linhas)


def _call_openrouter_extract_once(text: str, api_key: str, model: str = None,
                                   filename_hint: str = None, max_retries: int = 3,
                                   max_chars: int = 50000, pre_filtrar: bool = True,
                                   lista_referencia: list[dict] | None = None) -> dict:
    model = model or DEFAULT_EXTRACTION_MODEL
    url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }

    texto_base = pre_filtrar_texto(text) if pre_filtrar else text
    texto_truncado = len(texto_base) > max_chars
    user_content = texto_base[:max_chars]
    try:
        bloco_referencia = _formatar_lista_referencia(lista_referencia)
    except Exception:
        # lista de referência malformada nunca deve impedir a extração em si
        bloco_referencia = ""
    if bloco_referencia:
        user_content = f"{bloco_referencia}\n\n{user_content}"
    if filename_hint:
        user_content = f"Nome do arquivo: {filename_hint}\n\n{user_content}"

    def _montar_payload(com_schema: bool) -> dict:
        p = {
            "model": model,
            "messages": [
                {"role": "system", "content": EXTRACTION_SYSTEM_PROMPT},
                {"role": "user", "content": user_content},
            ],
            "temperature": 0,
            # OpenRouter devolve o custo real da chamada em usage.cost
            "usage": {"include": True},
        }
        if com_schema:
            p["response_format"] = {"type": "json_schema", "json_schema": EXTRACTION_JSON_SCHEMA}
        return p

    usar_schema = True
    last_error = "erro desconhecido"
    for attempt in range(max_retries):
        try:
            resp = SESSAO.post(url, headers=headers, json=_montar_payload(usar_schema), timeout=(10, 120))
            if resp.status_code in (400, 404) and usar_schema:
                # Provedor/modelo sem suporte a structured outputs: refaz sem schema
                usar_schema = False
                resp = SESSAO.post(url, headers=headers, json=_montar_payload(False), timeout=(10, 120))
            resp.raise_for_status()
            data = resp.json()
            content = data["choices"][0]["message"]["content"]
            content_clean = re.sub(r"^```json\s*|\s*```$", "", content.strip(), flags=re.MULTILINE)
            content_clean = content_clean.strip().strip("`")
            parsed = json.loads(content_clean)
            if not isinstance(parsed, dict):
                raise ValueError(
                    f"Resposta JSON da IA não é um objeto (tipo recebido: {type(parsed).__name__})"
                )
            parsed = _enrich_item_fields(parsed, texto_base)
            if isinstance(parsed.get("itens"), list):
                parsed["itens"] = [normalizar_item(i) for i in parsed["itens"] if isinstance(i, dict)]
            parsed["texto_truncado"] = texto_truncado
            parsed["usage"] = _extract_usage_info(data, model)
            return parsed
        except (requests.RequestException, json.JSONDecodeError, KeyError, IndexError, TypeError, ValueError) as exc:
            last_error = str(exc)
            if attempt < max_retries - 1:
                time.sleep(2 ** attempt)

    return {
        "empresa": filename_hint,
        "itens": [],
        "erro": last_error,
        "texto_truncado": texto_truncado,
        "usage": {
            "prompt_tokens": 0,
            "completion_tokens": 0,
            "total_tokens": 0,
            "cost_usd": 0.0,
            "estimated": True,
        },
    }


def call_openrouter_extract(text: str, api_key: str, model: str = None,
                             filename_hint: str = None, max_retries: int = 3,
                             max_chars: int = 50000, pre_filtrar: bool = True,
                             escalate: bool = True,
                             lista_referencia: list[dict] | None = None) -> dict:
    """
    Extração via LLM com escalonamento automático: tenta o modelo barato;
    se a extração falhar ou vier vazia, reextrai UMA vez com o modelo forte
    (ESCALATION_MODEL). O custo das duas chamadas é somado no usage.

    lista_referencia: itens oficiais da solicitação de orçamento/edital
    (numero_item, codigo, descricao), usados como contexto para a IA casar
    corretamente itens de fornecedores com nomenclatura diferente.
    """
    model = model or DEFAULT_EXTRACTION_MODEL
    parsed = _call_openrouter_extract_once(
        text, api_key, model=model, filename_hint=filename_hint,
        max_retries=max_retries, max_chars=max_chars, pre_filtrar=pre_filtrar,
        lista_referencia=lista_referencia,
    )

    extracao_ruim = bool(parsed.get("erro")) or not parsed.get("itens")
    if escalate and extracao_ruim and model != ESCALATION_MODEL:
        usage_barato = parsed.get("usage") or {}
        parsed_forte = _call_openrouter_extract_once(
            text, api_key, model=ESCALATION_MODEL, filename_hint=filename_hint,
            max_retries=max_retries, max_chars=max_chars, pre_filtrar=pre_filtrar,
            lista_referencia=lista_referencia,
        )
        if parsed_forte.get("itens") or not parsed.get("itens"):
            # usa o resultado do modelo forte; soma o gasto da tentativa barata
            usage_forte = parsed_forte.get("usage") or {}
            for campo in ("prompt_tokens", "completion_tokens", "total_tokens"):
                usage_forte[campo] = _to_int(usage_forte.get(campo)) + _to_int(usage_barato.get(campo))
            try:
                usage_forte["cost_usd"] = float(usage_forte.get("cost_usd") or 0) + float(usage_barato.get("cost_usd") or 0)
            except (TypeError, ValueError):
                pass
            parsed_forte["usage"] = usage_forte
            parsed_forte["escalado_para"] = ESCALATION_MODEL
            return parsed_forte

    return parsed


def _render_pdf_paginas_png(path: str, max_paginas: int = 8, resolution: int = 200) -> list[bytes]:
    """Renderiza as primeiras paginas do PDF como PNG (para extracao por visao)."""
    import io as _io
    pdfplumber = _import_pdfplumber()
    imagens = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages[:max_paginas]:
            try:
                img = page.to_image(resolution=resolution).original
                buf = _io.BytesIO()
                img.save(buf, format="PNG")
                imagens.append(buf.getvalue())
            except Exception:
                continue
    return imagens


def call_openrouter_vision_extract(imagens: list[bytes], api_key: str, model: str = None,
                                    filename_hint: str = None, max_retries: int = 2,
                                    mime: str = "image/png",
                                    lista_referencia: list[dict] | None = None) -> dict:
    """
    Extracao de itens diretamente da IMAGEM do documento, via modelo multimodal.

    Ultimo recurso da cascata de OCR: usado quando o Tesseract nao esta
    disponivel ou devolveu texto inutilizavel. O modelo de visao le layout,
    tabelas, carimbos e manuscrito melhor que OCR classico em documento ruim,
    ao custo de alguns centavos por documento.
    """
    import base64

    model = model or DEFAULT_EXTRACTION_MODEL
    url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}

    try:
        bloco_referencia = _formatar_lista_referencia(lista_referencia)
    except Exception:
        bloco_referencia = ""
    conteudo = [{
        "type": "text",
        "text": (
            (f"Nome do arquivo: {filename_hint}\n" if filename_hint else "")
            + (f"{bloco_referencia}\n\n" if bloco_referencia else "")
            + "Extraia os itens de orcamento das imagens deste documento, "
              "seguindo exatamente as regras do sistema."
        ),
    }]
    for img_bytes in imagens[:8]:
        b64 = base64.b64encode(img_bytes).decode("ascii")
        conteudo.append({
            "type": "image_url",
            "image_url": {"url": f"data:{mime};base64,{b64}"},
        })

    def _payload(com_schema: bool) -> dict:
        p = {
            "model": model,
            "messages": [
                {"role": "system", "content": EXTRACTION_SYSTEM_PROMPT},
                {"role": "user", "content": conteudo},
            ],
            "temperature": 0,
            "usage": {"include": True},
        }
        if com_schema:
            p["response_format"] = {"type": "json_schema", "json_schema": EXTRACTION_JSON_SCHEMA}
        return p

    usar_schema = True
    last_error = "erro desconhecido"
    for attempt in range(max_retries + 1):
        try:
            resp = SESSAO.post(url, headers=headers, json=_payload(usar_schema), timeout=(10, 180))
            if resp.status_code in (400, 404) and usar_schema:
                usar_schema = False
                resp = SESSAO.post(url, headers=headers, json=_payload(False), timeout=(10, 180))
            resp.raise_for_status()
            data = resp.json()
            content = data["choices"][0]["message"]["content"]
            content_clean = re.sub(r"^```json\s*|\s*```$", "", content.strip(), flags=re.MULTILINE)
            parsed = json.loads(content_clean.strip().strip("`"))
            if isinstance(parsed.get("itens"), list):
                parsed["itens"] = [normalizar_item(i) for i in parsed["itens"] if isinstance(i, dict)]
            parsed["texto_truncado"] = False
            parsed["usage"] = _extract_usage_info(data, model)
            parsed["fonte_extracao"] = "visao"
            return parsed
        except (requests.RequestException, json.JSONDecodeError, KeyError, IndexError) as exc:
            last_error = str(exc)
            if attempt < max_retries:
                time.sleep(2 ** attempt)

    return {
        "empresa": filename_hint,
        "itens": [],
        "erro": f"Extracao por visao falhou: {last_error}",
        "texto_truncado": False,
        "usage": {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0,
                   "cost_usd": 0.0, "estimated": True},
    }


def _extract_text_pages_pdf(path: str, ocr_lang: str = "por"):
    chave = _chave_arquivo(path)
    with _PAGINAS_CACHE_LOCK:
        if chave in _PAGINAS_CACHE:
            return _PAGINAS_CACHE[chave]

    pdfplumber = _import_pdfplumber()
    pages = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = (page.extract_text() or "").strip()
            veio_ocr = False
            erro_ocr = None
            tentou_ocr = _texto_pobre_para_ocr(page_text)
            if tentou_ocr:
                try:
                    page_img = page.to_image(resolution=250).original
                    ocr_text, erro_ocr = _ocr_image(page_img, ocr_lang=ocr_lang)
                    if _preferir_texto_ocr(page_text, ocr_text):
                        page_text = ocr_text
                        veio_ocr = True
                except Exception as exc:
                    erro_ocr = str(exc)
            pages.append({
                "texto": page_text or "",
                "veio_ocr": veio_ocr,
                "tentou_ocr": tentou_ocr,
                "erro_ocr": erro_ocr,
            })
    with _PAGINAS_CACHE_LOCK:
        _PAGINAS_CACHE[chave] = pages
    return pages


def _inferir_origem_pagina(item: dict, pages: list[dict]) -> str:
    descricao = limpar_quebras_e_caracteres(str(item.get("descricao") or "")).lower()
    numero = str(item.get("numero_item") or "").strip()
    if not pages:
        return "documento"

    for idx, p in enumerate(pages, start=1):
        texto = p.get("texto", "").lower()
        if numero and re.search(rf"\b{re.escape(numero)}\b", texto):
            return f"Pagina {idx}"
        if descricao and descricao[:40] and descricao[:40] in texto:
            return f"Pagina {idx}"

    return "Pagina 1"


def _anotar_fonte_origem(itens: list[dict], fonte: str, pages: list[dict] | None = None) -> list[dict]:
    saida = []
    for item in itens or []:
        i2 = dict(item)
        i2["fonte_extracao"] = fonte
        if not i2.get("origem"):
            i2["origem"] = _inferir_origem_pagina(i2, pages or [])
        saida.append(normalizar_item(i2))
    return saida


def calcular_similaridade_item(item_parser: dict, item_ia: dict) -> float:
    num_parser = str(item_parser.get("numero_item") or "").strip()
    num_ia = str(item_ia.get("numero_item") or "").strip()
    if num_parser and num_ia and num_parser == num_ia:
        return 100.0

    desc_parser = limpar_quebras_e_caracteres(str(item_parser.get("descricao") or "")).lower()
    desc_ia = limpar_quebras_e_caracteres(str(item_ia.get("descricao") or "")).lower()
    sim_desc = token_set_ratio(desc_parser, desc_ia)

    unidade_parser = str(item_parser.get("unidade") or "").upper()
    unidade_ia = str(item_ia.get("unidade") or "").upper()
    sim_unidade = 100.0 if unidade_parser and unidade_parser == unidade_ia else 0.0

    qtd_p = item_parser.get("quantidade")
    qtd_i = item_ia.get("quantidade")
    sim_qtd = 0.0
    if qtd_p is not None and qtd_i is not None:
        if qtd_p == qtd_i:
            sim_qtd = 100.0
        elif max(abs(qtd_p), abs(qtd_i)) > 0:
            diff = abs(qtd_p - qtd_i) / max(abs(qtd_p), abs(qtd_i))
            if diff < 0.05:
                sim_qtd = 100.0

    return (0.60 * sim_desc) + (0.25 * sim_unidade) + (0.15 * sim_qtd)


def comparar_extracoes(itens_parser: list[dict], itens_ia: list[dict]):
    def _mesclar_item(parser_item: dict, ia_item: dict, score: float, numerico_bate: bool) -> dict:
        item_final = dict(parser_item)
        item_final["fonte_extracao"] = "dupla_checagem"
        item_final["confianca"] = "alta" if score >= 70 and numerico_bate else "baixa"

        if ia_item.get("descricao") and len(str(ia_item.get("descricao"))) > len(str(item_final.get("descricao") or "")):
            item_final["descricao"] = ia_item.get("descricao")

        for campo in ("unidade", "quantidade", "preco_total", "origem"):
            if item_final.get(campo) in (None, "") and ia_item.get(campo) not in (None, ""):
                item_final[campo] = ia_item.get(campo)

        if item_final.get("preco_unitario") is None and ia_item.get("preco_unitario") is not None:
            item_final["preco_unitario"] = ia_item.get("preco_unitario")

        return normalizar_item(item_final)

    pares = []
    for i, item_p in enumerate(itens_parser):
        for j, item_i in enumerate(itens_ia):
            score = calcular_similaridade_item(item_p, item_i)
            pares.append((score, i, j))

    pares.sort(reverse=True, key=lambda x: x[0])
    usados_parser = set()
    usados_ia = set()
    casamentos = []
    for score, i, j in pares:
        if i in usados_parser or j in usados_ia:
            continue
        usados_parser.add(i)
        usados_ia.add(j)
        casamentos.append((score, itens_parser[i], itens_ia[j]))

    itens_finais = []
    conflitos = []
    for score, parser_item, ia_item in casamentos:
        p_unit = parser_item.get("preco_unitario")
        i_unit = ia_item.get("preco_unitario")
        numerico_bate = False
        if p_unit is not None and i_unit is not None:
            maior_valor = max(abs(p_unit), abs(i_unit))
            if maior_valor == 0:
                numerico_bate = True
            else:
                numerico_bate = abs(p_unit - i_unit) / maior_valor < 0.01

        itens_finais.append(_mesclar_item(parser_item, ia_item, score, numerico_bate))

        if score < 70 or not numerico_bate:
            conflitos.append({
                "tipo": "divergência parser vs. IA",
                "numero_item": parser_item.get("numero_item") or ia_item.get("numero_item"),
                "descricao_nova": parser_item.get("descricao") or "",
                "casou_com": ia_item.get("descricao") or "",
                "score": round(score, 2),
            })

    for idx, parser_item in enumerate(itens_parser):
        if idx not in usados_parser:
            item_final = dict(parser_item)
            item_final["fonte_extracao"] = "dupla_checagem"
            item_final["confianca"] = "baixa"
            itens_finais.append(normalizar_item(item_final))
            conflitos.append({
                "tipo": "divergência parser vs. IA",
                "numero_item": parser_item.get("numero_item"),
                "descricao_nova": parser_item.get("descricao") or "",
                "casou_com": "sem par na IA",
                "score": 0,
            })

    for idx, ia_item in enumerate(itens_ia):
        if idx not in usados_ia:
            item_final = dict(ia_item)
            item_final["fonte_extracao"] = "dupla_checagem"
            item_final["confianca"] = "média"
            itens_finais.append(normalizar_item(item_final))
            conflitos.append({
                "tipo": "divergência parser vs. IA",
                "numero_item": ia_item.get("numero_item"),
                "descricao_nova": ia_item.get("descricao") or "",
                "casou_com": "sem par no parser",
                "score": 0,
            })

    return itens_finais, conflitos


# Domínios de e-mail genéricos: nunca usar como nome de empresa (pessoa física
# ou provedor gratuito não identifica o fornecedor).
_DOMINIOS_EMAIL_GENERICOS = {
    "gmail.com", "hotmail.com", "outlook.com", "outlook.com.br", "yahoo.com",
    "yahoo.com.br", "bol.com.br", "uol.com.br", "terra.com.br", "live.com",
    "icloud.com", "globo.com", "ig.com.br",
}

# Se a razão social "detectada" contiver um destes termos, é o comprador
# (Marinha), não o fornecedor — descarta o resultado.
_TERMOS_COMPRADOR = ("MARINHA", "COMANDO", "COMRJ")


def _nome_from_header(from_hdr: str) -> str | None:
    """Extrai o nome de exibição (ou usuário do e-mail) do cabeçalho From."""
    if not from_hdr:
        return None
    m = re.match(r"(?P<name>.+?)\s*<.+?>", from_hdr)
    if m:
        nome = m.group("name").strip().strip('"')
        return nome or None
    nome = (from_hdr.split("@")[0] if "@" in from_hdr else from_hdr).strip()
    return nome or None


def _dominio_email_para_nome(email_remetente: str | None) -> str | None:
    """Deriva um nome de empresa a partir do domínio do e-mail, descartando
    provedores genéricos (gmail, hotmail etc.)."""
    if not email_remetente or "@" not in email_remetente:
        return None
    dominio = email_remetente.strip().lower().split("@")[-1].strip(">").strip()
    if not dominio or dominio in _DOMINIOS_EMAIL_GENERICOS:
        return None
    primeiro_rotulo = dominio.split(".")[0]
    if not primeiro_rotulo:
        return None
    return primeiro_rotulo.upper()


def _extrair_email_de_header(from_hdr: str) -> str | None:
    if not from_hdr:
        return None
    m = re.search(r"<(?P<email>[^<>]+@[^<>]+)>", from_hdr)
    if m:
        return m.group("email").strip()
    if "@" in from_hdr:
        return from_hdr.strip()
    return None


def _extrair_razao_social_xlsx_aba2(path: str) -> str | None:
    """Lê o campo 'RAZÃO SOCIAL' da aba 'Aba2 MODELO DE PROPOSTA' (ou
    equivalente) de um .xlsx: rótulo na primeira coluna, valor em alguma
    coluna à direita, na mesma linha, dentro das primeiras ~18 linhas."""
    if not path.lower().endswith((".xlsx", ".xls")):
        return None
    try:
        openpyxl = _import_openpyxl()
        wb = openpyxl.load_workbook(path, data_only=True)
    except Exception:
        return None

    def _norm(txt) -> str:
        if txt is None:
            return ""
        base = unicodedata.normalize("NFKD", str(txt))
        sem_acento = "".join(ch for ch in base if not unicodedata.combining(ch))
        return re.sub(r"\s+", " ", sem_acento).strip().upper()

    sheets = list(wb.worksheets)
    # Prioriza a aba cujo nome sugere ser o modelo de proposta.
    sheets.sort(key=lambda s: 0 if ("MODELO" in _norm(s.title) and "PROPOSTA" in _norm(s.title)) else 1)

    for sheet in sheets:
        for row in sheet.iter_rows(min_row=1, max_row=18, values_only=True):
            if not row:
                continue
            rotulo = _norm(row[0])
            if "RAZAO SOCIAL" not in rotulo:
                continue
            for valor in row[1:]:
                if valor is None:
                    continue
                valor_str = str(valor).strip()
                if not valor_str or valor_str == "*":
                    continue
                return valor_str
    return None


def resolver_nome_empresa(
    *,
    xlsx_paths: list[str] | None = None,
    texto: str | None = None,
    remetente_email: str | None = None,
    from_hdr: str | None = None,
) -> str | None:
    """Resolve o nome do fornecedor com uma única ordem de confiança:
    1) RAZÃO SOCIAL da aba modelo de proposta do(s) xlsx anexado(s);
    2) razão social próxima ao CNPJ no texto (descartando o comprador);
    3) domínio do e-mail do remetente (descartando provedores genéricos);
    4) nome de exibição do cabeçalho From.
    """
    for xlsx_path in (xlsx_paths or []):
        nome = _extrair_razao_social_xlsx_aba2(xlsx_path)
        if nome:
            return nome

    if texto:
        candidata = extrair_razao_social(texto)
        if candidata:
            candidata_up = candidata.upper()
            if not any(termo in candidata_up for termo in _TERMOS_COMPRADOR):
                return candidata

    email_efetivo = remetente_email or _extrair_email_de_header(from_hdr or "")
    nome_dominio = _dominio_email_para_nome(email_efetivo)
    if nome_dominio:
        return nome_dominio

    return _nome_from_header(from_hdr or "")


_PALAVRAS_NAO_IDENTIFICADORAS_EMPRESA = {
    "A", "COMERCIO", "COMERCIAL", "COMPANHIA", "CONSTRUCOES", "DE", "DA", "DO",
    "DADOS", "DAS", "DOS", "E", "EIRELI", "EMPRESA", "ENDERECO", "ENGENHARIA",
    "EPP", "FORNECEDOR", "IDENTIFICACAO", "IRELI", "LTDA", "ME", "MEI", "ORCAMENTO",
    "PROPOSTA", "RAZAO", "SA", "SERVICO", "SERVICOS", "SOCIAL", "SOCIEDADE", "SOLUCOES",
    "UN", "UND", "UNIDADE",
}


def normalizar_nome_empresa_curto(nome: str | None) -> str | None:
    """Retorna uma marca curta e legivel para o cabecalho do mapa comparativo."""
    bruto = (nome or "").strip()
    if not bruto:
        return None

    emails = re.findall(r"[\w.+-]+@[\w.-]+", bruto)
    if emails:
        dominio = emails[0].split("@", 1)[1].split(".")[0]
        dominio = unicodedata.normalize("NFKD", dominio)
        dominio = "".join(ch for ch in dominio if not unicodedata.combining(ch)).upper()
        for sufixo in ("COMERCIO", "ENGENHARIA", "INDUSTRIA", "SERVICOS", "SOLUCOES"):
            if dominio.endswith(sufixo) and len(dominio) - len(sufixo) >= 3:
                dominio = dominio[:-len(sufixo)]
                break
        if re.fullmatch(r"[A-Z]{3,}", dominio):
            return dominio

    texto = unicodedata.normalize("NFKD", bruto)
    texto = "".join(ch for ch in texto if not unicodedata.combining(ch))
    texto = re.sub(r"\S+@\S+", " ", texto)
    texto = re.sub(r"\b\d+[\w./-]*\b", " ", texto)
    tokens = re.findall(r"[A-Za-z]{2,}", texto.upper())
    for token in tokens:
        if token not in _PALAVRAS_NAO_IDENTIFICADORAS_EMPRESA:
            return token
    return None


def _score_ancoragem(itens: list[dict]) -> int:
    """Conta itens com número de item ou código de estoque (PI) preenchido —
    sinal de que o anexo traz a referência oficial do edital."""
    score = 0
    for item in itens or []:
        numero_item = item.get("numero_item")
        if numero_item not in (None, "", "None"):
            score += 1
            continue
        codigo = item.get("codigo")
        if codigo and re.fullmatch(r"\d{5,}", str(codigo).strip()):
            score += 1
    return score


def _chave_item(item: dict) -> str:
    numero_item = item.get("numero_item")
    if numero_item not in (None, "", "None"):
        return f"num:{str(numero_item).strip().lower()}"
    descricao = (item.get("descricao") or "").strip().lower()
    return f"desc:{descricao}"


def _mesclar_resultados_anexos(resultados: list[dict]) -> dict:
    """Escolhe como base o anexo com maior 'ancoragem' (mais itens com número
    de item/código de estoque) e completa com itens exclusivos dos demais
    anexos, sem sobrescrever nada do resultado base."""
    if len(resultados) == 1:
        return resultados[0]

    scores = [_score_ancoragem(r.get("itens", [])) for r in resultados]
    idx_base = max(range(len(resultados)), key=lambda i: scores[i])
    base = dict(resultados[idx_base])
    itens_base = list(base.get("itens", []))
    chaves_base = {_chave_item(i) for i in itens_base}

    itens_extras = []
    review_extra = []
    usage_extra = {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0, "cost_usd": 0.0}
    debug_extra = []
    truncado_extra = False

    def _mesmo_item_cotado(base_item: dict, complemento: dict) -> bool:
        """Reconhece o mesmo item entre PDF comercial e planilha do edital.

        O número/código pode mudar entre os documentos, mas quantidade, preço
        unitário e total formam uma assinatura forte da cotação preenchida.
        """
        try:
            campos = ("quantidade", "preco_unitario", "preco_total")
            for campo in campos:
                valor_base = float(base_item.get(campo))
                valor_complemento = float(complemento.get(campo))
                tolerancia = max(0.01, 0.001 * abs(valor_base))
                if abs(valor_base - valor_complemento) > tolerancia:
                    return False
            return True
        except (TypeError, ValueError):
            return False

    for i, res in enumerate(resultados):
        if i == idx_base:
            continue
        for item in res.get("itens", []):
            chave = _chave_item(item)
            if chave in chaves_base:
                continue
            item_base = next((existente for existente in itens_base if _mesmo_item_cotado(existente, item)), None)
            if item_base is not None:
                descricao_complementar = str(item.get("descricao") or "")
                if len(descricao_complementar) > len(str(item_base.get("descricao") or "")):
                    item_base["descricao"] = descricao_complementar
                continue
            chaves_base.add(chave)
            itens_extras.append(item)
        review_extra.extend(res.get("review", []) or [])
        usage_res = res.get("usage") or {}
        for campo in ("prompt_tokens", "completion_tokens", "total_tokens", "cost_usd"):
            try:
                usage_extra[campo] += float(usage_res.get(campo) or 0)
            except (TypeError, ValueError):
                pass
        debug_extra.append(
            f"Anexo complementar '{os.path.basename(str(res.get('_path', '')))}': "
            f"{len(res.get('itens', []))} item(ns) considerado(s), "
            f"{len(itens_extras)} incorporado(s) ao resultado final"
        )
        truncado_extra = truncado_extra or bool(res.get("texto_truncado"))

    base["itens"] = itens_base + itens_extras
    base["review"] = list(base.get("review", []) or []) + review_extra
    base["texto_truncado"] = bool(base.get("texto_truncado")) or truncado_extra
    base_usage = dict(base.get("usage") or {})
    for campo in ("prompt_tokens", "completion_tokens", "total_tokens"):
        base_usage[campo] = int(base_usage.get(campo, 0) or 0) + int(usage_extra[campo])
    base_usage["cost_usd"] = float(base_usage.get("cost_usd", 0) or 0) + usage_extra["cost_usd"]
    base_usage["estimated"] = bool(base_usage.get("estimated")) or any(
        (r.get("usage") or {}).get("estimated") for r in resultados
    )
    base["usage"] = base_usage
    base["debug_events"] = list(base.get("debug_events", []) or []) + debug_extra
    if len(resultados) > 1:
        base.setdefault("debug_highlights", []).append(
            f"RELEVANTE: {len(resultados)} anexo(s) processado(s); anexo base escolhido por ter "
            f"mais itens com número de item/código de estoque (ancoragem={scores[idx_base]})"
        )
    return base


def extrair_anexos_orcamento_em_camadas(paths: list[str], api_key: str, model: str,
                                        pre_filtrar: bool = True,
                                        limiar_alto: int = 85,
                                        limiar_baixo: int = 40,
                                        lista_referencia: list[dict] | None = None) -> dict:
    """Extrai e consolida anexos complementares de uma mesma resposta de fornecedor."""
    resultados = []
    for path in paths:
        try:
            resultado = extrair_orcamento_em_camadas(
                path, api_key, model, pre_filtrar, limiar_alto, limiar_baixo, lista_referencia
            )
            resultado["_path"] = path
            resultados.append(resultado)
        except Exception:
            continue
    if not resultados:
        return {"empresa": None, "itens": [], "erro": "Nenhum anexo pôde ser processado."}
    resultado = _mesclar_resultados_anexos(resultados)
    resultado.pop("_path", None)
    return resultado


def extrair_orcamento_em_camadas(path: str, api_key: str, model: str,
                                 pre_filtrar: bool = True,
                                 limiar_alto: int = 85,
                                 limiar_baixo: int = 40,
                                 lista_referencia: list[dict] | None = None) -> dict:
    """Pipeline: a IA sempre interpreta o conteudo; a extracao estrutural
    (quando disponivel) e usada como segunda fonte para conferencia (dupla
    checagem), nunca para pular a IA por completo. `lista_referencia` (itens
    oficiais da solicitacao de orcamento/edital) e passada a IA para ajudar
    a casar itens de fornecedores diferentes mesmo com nomenclatura distinta."""
    if os.path.splitext(path)[1].lower() == ".pdf":
        pages = _extract_text_pages_pdf(path)
        tipo = detectar_tipo_por_paginas(pages)
    else:
        tipo = detectar_tipo_e_rotear(path)
    nome_arquivo = os.path.basename(path)
    review = []
    confianca = None
    texto_base = ""
    texto_truncado = False
    fonte_global = "ia"
    debug_events = []
    debug_highlights = []

    debug_events.append(f"Arquivo recebido: {nome_arquivo}")
    debug_events.append(f"Percepcao inicial: tipo detectado = {tipo}")

    if tipo == "xlsx":
        debug_events.append("Biblioteca acionada: openpyxl (extracao estrutural de planilha)")
    elif tipo == "csv":
        debug_events.append("Biblioteca acionada: csv (leitura textual estruturada)")
    elif tipo == "docx":
        debug_events.append("Biblioteca acionada: python-docx (extracao estrutural de documento Word)")
    elif tipo in ("pdf_texto", "pdf_escaneado"):
        debug_events.append("Biblioteca acionada: pdfplumber (leitura de paginas PDF)")
        debug_events.append("Biblioteca potencial: requests (chamada IA quando necessario)")
    elif tipo == "imagem":
        debug_events.append("Biblioteca acionada: Pillow + pytesseract (OCR de imagem)")
        debug_events.append("Biblioteca potencial: requests (chamada IA quando necessario)")

    if tipo in ("xlsx", "docx"):
        itens_estrutural = extrair_xlsx_estruturado(path) if tipo == "xlsx" else extrair_docx_estruturado(path)
        if itens_estrutural:
            texto_base = extract_text(path)
            empresa_det = extrair_razao_social(texto_base)
            debug_events.append(f"Achado: {len(itens_estrutural)} item(ns) extraido(s) por modo estrutural")

            if api_key:
                llm = call_openrouter_extract(
                    texto_base, api_key, model=model, filename_hint=nome_arquivo,
                    pre_filtrar=pre_filtrar, lista_referencia=lista_referencia,
                )
                itens_ia = _anotar_fonte_origem(llm.get("itens", []), "ia")
                itens_finais, conflitos = comparar_extracoes(itens_estrutural, itens_ia)
                review.extend(conflitos)
                fonte_global = "dupla_checagem"
                empresa_final = llm.get("empresa") or empresa_det or os.path.splitext(nome_arquivo)[0]
                usage = llm.get("usage") or {"prompt_tokens": 0, "completion_tokens": 0,
                                              "total_tokens": 0, "cost_usd": 0.0, "estimated": True}
                debug_events.append(f"Achado IA: {len(itens_ia)} item(ns) retornado(s), mesclado com o estrutural")
                debug_highlights.append(
                    "RELEVANTE: IA conferiu a extracao estrutural (casamento com lista de referencia quando houver)"
                )
            else:
                itens_finais = itens_estrutural
                fonte_global = "estrutural"
                empresa_final = empresa_det or os.path.splitext(nome_arquivo)[0]
                usage = {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0,
                         "cost_usd": 0.0, "estimated": False}
                debug_highlights.append("RELEVANTE: sem chave de IA configurada, usado apenas o estrutural")

            return {
                "empresa": empresa_final,
                "itens": itens_finais,
                "review": review,
                "confianca_estrutural": 100,
                "fonte_processamento": fonte_global,
                "texto_truncado": False,
                "debug_events": debug_events,
                "debug_highlights": debug_highlights,
                "cnpj": extrair_cnpj(texto_base),
                "telefone": extrair_telefone(texto_base),
                "usage": usage,
            }

        if tipo == "eml":
            # processa .eml: se houver anexos, delega ao anexo; se nao, usa corpo do email
            try:
                import email
                with open(path, 'rb') as fh:
                    raw = fh.read()
                msg = email.message_from_bytes(raw)
            except Exception:
                # nao conseguiu abrir como eml: cair no fluxo normal
                msg = None
            if msg is not None:
                # coleta anexos
                attach_paths = []
                for idx, part in enumerate(msg.walk()):
                    fn = part.get_filename()
                    if fn:
                        payload = part.get_payload(decode=True)
                        if payload:
                            tmp = f"/tmp/{nome_arquivo}.attach.{idx}.{fn}"
                            try:
                                with open(tmp, 'wb') as out:
                                    out.write(payload)
                                attach_paths.append(tmp)
                            except Exception:
                                continue
                if attach_paths:
                    from_hdr = msg.get('From') or ''
                    remetente_email = _extrair_email_de_header(from_hdr)

                    # processa TODOS os anexos individualmente; erro em um nao aborta os demais
                    resultados_attach = []
                    for attach_path in attach_paths:
                        try:
                            res = extrair_orcamento_em_camadas(
                                attach_path, api_key, model, pre_filtrar, limiar_alto, limiar_baixo
                            )
                            res['_path'] = attach_path
                            resultados_attach.append(res)
                        except Exception as exc:
                            debug_events.append(
                                f"Falha ao processar anexo '{os.path.basename(attach_path)}': {exc}"
                            )
                            continue

                    if not resultados_attach:
                        # todos os anexos falharam: cai no fluxo sem anexos (usa corpo do email)
                        attach_paths = []
                    else:
                        res_attach = _mesclar_resultados_anexos(resultados_attach)
                        res_attach.pop('_path', None)

                        xlsx_paths = [p for p in attach_paths if p.lower().endswith((".xlsx", ".xls"))]
                        textos_anexos = []
                        for p in attach_paths:
                            try:
                                textos_anexos.append(extract_text(p))
                            except Exception:
                                continue
                        texto_para_empresa = "\n".join(textos_anexos) if textos_anexos else None
                        empresa_resolvida = resolver_nome_empresa(
                            xlsx_paths=xlsx_paths,
                            texto=texto_para_empresa,
                            remetente_email=remetente_email,
                            from_hdr=from_hdr,
                        )
                        if empresa_resolvida:
                            res_attach['empresa'] = empresa_resolvida
                        return res_attach

                # sem anexos: monta texto a partir das partes textuais
                parts = []
                for part in msg.walk():
                    if part.get_content_type() == 'text/plain' and not part.get_filename():
                        try:
                            payload = part.get_payload(decode=True)
                            if payload:
                                parts.append(payload.decode('utf-8', errors='ignore'))
                        except Exception:
                            pass
                texto_base = "\n".join(parts) if parts else ''
                pages = [{"texto": texto_base, "veio_ocr": False, "tentou_ocr": False, "erro_ocr": None}]

                llm = call_openrouter_extract(
                    texto_base,
                    api_key,
                    model=model,
                    filename_hint=nome_arquivo,
                    pre_filtrar=pre_filtrar,
                )
                itens_ia = _anotar_fonte_origem(llm.get("itens", []), "ia", pages)
                itens_finais = itens_ia

                from_hdr = msg.get('From') or ''
                empresa_resolvida = resolver_nome_empresa(
                    texto=texto_base,
                    remetente_email=_extrair_email_de_header(from_hdr),
                    from_hdr=from_hdr,
                )
                empresa_llm = empresa_resolvida or llm.get("empresa")
                empresa_det = empresa_resolvida

                # Corrige preco_unitario quando necessario
                itens_corrigidos = []
                for i in itens_finais:
                    p_unit = i.get('preco_unitario')
                    p_total = i.get('preco_total')
                    qtd = i.get('quantidade')
                    try:
                        if (p_unit is None or (p_total is not None and qtd)) and p_total is not None and qtd:
                            i['preco_unitario'] = float(p_total) / float(qtd)
                        elif p_unit is not None and p_total is not None and qtd:
                            if float(p_unit) > float(p_total):
                                i['preco_unitario'] = float(p_total) / float(qtd)
                    except Exception:
                        pass
                    itens_corrigidos.append(i)

                return {
                    'empresa': empresa_llm or empresa_det or os.path.splitext(nome_arquivo)[0],
                    'itens': [normalizar_item(i) for i in itens_corrigidos],
                    'review': review,
                    'confianca_estrutural': None,
                    'fonte_processamento': 'ia',
                    'texto_truncado': bool(llm.get('texto_truncado')),
                    'cnpj': extrair_cnpj(texto_base),
                    'telefone': extrair_telefone(texto_base),
                    'debug_events': debug_events,
                    'debug_highlights': debug_highlights,
                    'usage': llm.get('usage', {}),
                }
    if tipo in ("pdf_texto", "pdf_escaneado"):
        pages = _extract_text_pages_pdf(path)
        texto_base = "\n".join(p["texto"] for p in pages)
        debug_events.append(f"Achado: {len(pages)} pagina(s) lida(s) no PDF")
        debug_events.append(f"Achado: {len(texto_base)} caractere(s) extraido(s) do documento")
        paginas_ocr = sum(1 for p in pages if p.get("veio_ocr"))
        paginas_tentou_ocr = sum(1 for p in pages if p.get("tentou_ocr"))
        if paginas_tentou_ocr:
            debug_events.append(f"Achado: OCR tentado em {paginas_tentou_ocr} pagina(s)")
        if paginas_ocr:
            debug_events.append(f"Achado: OCR aplicado em {paginas_ocr} pagina(s)")

        if not texto_base.strip():
            # --- Fallback de visao: OCR falhou/indisponivel -> manda a IMAGEM
            # das paginas para um modelo multimodal (ultimo recurso da cascata)
            if api_key:
                try:
                    imagens = _render_pdf_paginas_png(path)
                except Exception:
                    imagens = []
                if imagens:
                    debug_events.append(
                        f"Fallback de visao: {len(imagens)} pagina(s) enviadas como imagem ao modelo multimodal"
                    )
                    llm_visao = call_openrouter_vision_extract(
                        imagens, api_key, model=model, filename_hint=nome_arquivo,
                        lista_referencia=lista_referencia,
                    )
                    if llm_visao.get("itens"):
                        debug_highlights.append(
                            "RELEVANTE: itens extraidos por VISAO (OCR classico falhou neste documento)"
                        )
                        return {
                            "empresa": llm_visao.get("empresa") or os.path.splitext(nome_arquivo)[0],
                            "itens": [normalizar_item(i) for i in llm_visao["itens"]],
                            "review": review,
                            "confianca_estrutural": None,
                            "fonte_processamento": "visao",
                            "texto_truncado": False,
                            "cnpj": None,
                            "telefone": None,
                            "debug_events": debug_events,
                            "debug_highlights": debug_highlights,
                            "usage": llm_visao.get("usage", {}),
                        }

            erros_ocr = [p.get("erro_ocr") for p in pages if p.get("erro_ocr")]
            mensagem_erro = "Texto vazio apos leitura do PDF; nao foi possivel extrair conteudo util."
            if erros_ocr:
                mensagem_erro = (
                    "Texto vazio apos leitura do PDF e OCR indisponivel/falhou. "
                    f"Detalhe: {erros_ocr[0]}"
                )
                debug_highlights.append("RELEVANTE: OCR nao disponivel no ambiente")
            return {
                "empresa": os.path.splitext(nome_arquivo)[0],
                "itens": [],
                "erro": mensagem_erro,
                "review": review,
                "confianca_estrutural": 0,
                "fonte_processamento": "ia",
                "texto_truncado": False,
                "cnpj": None,
                "telefone": None,
                "debug_events": debug_events,
                "debug_highlights": debug_highlights,
                "usage": {
                    "prompt_tokens": 0,
                    "completion_tokens": 0,
                    "total_tokens": 0,
                    "cost_usd": 0.0,
                    "estimated": False,
                },
            }

        resultado_estrutural = tentar_extracao_estrutural_pdf(path)
        confianca = calcular_confianca_estrutural(resultado_estrutural)
        if tipo == "pdf_escaneado":
            confianca = max(0, confianca - 10)
            debug_events.append("Ajuste aplicado: PDF escaneado detectado, penalidade de confianca estrutural")

        review.extend(resultado_estrutural.get("review") or [])
        debug_events.append(f"Percepcao do sistema: confianca estrutural = {confianca}")

        llm = call_openrouter_extract(
            texto_base,
            api_key,
            model=model,
            filename_hint=nome_arquivo,
            pre_filtrar=pre_filtrar,
            lista_referencia=lista_referencia,
        )
        debug_events.append(f"Biblioteca acionada: requests (modelo IA = {model})")
        texto_truncado = bool(llm.get("texto_truncado"))
        if texto_truncado:
            debug_highlights.append("RELEVANTE: texto truncado para caber no limite enviado a IA")
        itens_ia = _anotar_fonte_origem(llm.get("itens", []), "ia", pages)
        debug_events.append(f"Achado IA: {len(itens_ia)} item(ns) retornado(s)")

        if confianca is not None and confianca < limiar_baixo:
            fonte_global = "ia"
            itens_finais = itens_ia
            debug_highlights.append(
                f"RELEVANTE: confianca estrutural baixa ({confianca}) -> decisao final por IA"
            )
        elif resultado_estrutural.get("itens"):
            fonte_global = "dupla_checagem"
            itens_finais, conflitos = comparar_extracoes(resultado_estrutural["itens"], itens_ia)
            review.extend(conflitos)
            debug_events.append(
                f"Percepcao: dupla checagem executada (parser vs IA), conflitos={len(conflitos)}"
            )
            if confianca is not None and confianca >= limiar_alto:
                debug_highlights.append(
                    f"RELEVANTE: confianca estrutural alta ({confianca}), mas IA conferiu mesmo assim"
                )
            if conflitos:
                debug_highlights.append(
                    f"RELEVANTE: divergencias detectadas entre parser e IA ({len(conflitos)})"
                )
        else:
            fonte_global = "ia"
            itens_finais = itens_ia
            debug_events.append("Percepcao: sem base estrutural valida, decisao final por IA")

        empresa_det = extrair_razao_social(texto_base)
        empresa_llm = llm.get("empresa")
        if empresa_llm:
            debug_events.append(f"Achado: empresa identificada pela IA = {empresa_llm}")
        elif empresa_det:
            debug_events.append(f"Achado: empresa inferida por heuristica = {empresa_det}")

        if extrair_cnpj(texto_base):
            debug_events.append("Achado: CNPJ identificado no documento")

        return {
            "empresa": empresa_llm or empresa_det or os.path.splitext(nome_arquivo)[0],
            "itens": [normalizar_item(i) for i in itens_finais],
            "review": review,
            "confianca_estrutural": confianca,
            "fonte_processamento": fonte_global,
            "texto_truncado": texto_truncado,
            "cnpj": extrair_cnpj(texto_base),
            "telefone": extrair_telefone(texto_base),
            "debug_events": debug_events,
            "debug_highlights": debug_highlights,
            "usage": llm.get("usage", {
                "prompt_tokens": 0,
                "completion_tokens": 0,
                "total_tokens": 0,
                "cost_usd": 0.0,
                "estimated": True,
            }),
        }

    if tipo == "imagem":
        texto_base, erro_ocr = extract_text_from_image(path)
        debug_events.append(f"Achado: {len(texto_base)} caractere(s) extraido(s) da imagem")

        # --- Fallback de visao para imagem sem texto OCR utilizavel ----------
        if not texto_base.strip() and api_key:
            try:
                with open(path, "rb") as _fh:
                    img_bytes = _fh.read()
                ext_img = os.path.splitext(path)[1].lower().lstrip(".")
                mime_img = f"image/{'jpeg' if ext_img in ('jpg', 'jpeg') else ext_img or 'png'}"
                llm_visao = call_openrouter_vision_extract(
                    [img_bytes], api_key, model=model,
                    filename_hint=nome_arquivo, mime=mime_img,
                    lista_referencia=lista_referencia,
                )
            except Exception:
                llm_visao = {}
            if llm_visao.get("itens"):
                debug_events.append("Fallback de visao: imagem enviada ao modelo multimodal")
                debug_highlights.append(
                    "RELEVANTE: itens extraidos por VISAO (OCR classico falhou nesta imagem)"
                )
                return {
                    "empresa": llm_visao.get("empresa") or os.path.splitext(nome_arquivo)[0],
                    "itens": [normalizar_item(i) for i in llm_visao["itens"]],
                    "review": review,
                    "confianca_estrutural": None,
                    "fonte_processamento": "visao",
                    "texto_truncado": False,
                    "cnpj": None,
                    "telefone": None,
                    "debug_events": debug_events,
                    "debug_highlights": debug_highlights,
                    "usage": llm_visao.get("usage", {}),
                }

        if erro_ocr and not texto_base.strip():
            debug_highlights.append("RELEVANTE: OCR nao disponivel no ambiente")
            return {
                "empresa": os.path.splitext(nome_arquivo)[0],
                "itens": [],
                "erro": f"Falha no OCR da imagem: {erro_ocr}",
                "review": review,
                "confianca_estrutural": 0,
                "fonte_processamento": "ia",
                "texto_truncado": False,
                "cnpj": None,
                "telefone": None,
                "debug_events": debug_events,
                "debug_highlights": debug_highlights,
                "usage": {
                    "prompt_tokens": 0,
                    "completion_tokens": 0,
                    "total_tokens": 0,
                    "cost_usd": 0.0,
                    "estimated": False,
                },
            }

        if not texto_base.strip():
            return {
                "empresa": os.path.splitext(nome_arquivo)[0],
                "itens": [],
                "erro": "A imagem nao possui texto legivel para extracao.",
                "review": review,
                "confianca_estrutural": 0,
                "fonte_processamento": "ia",
                "texto_truncado": False,
                "cnpj": None,
                "telefone": None,
                "debug_events": debug_events,
                "debug_highlights": debug_highlights,
                "usage": {
                    "prompt_tokens": 0,
                    "completion_tokens": 0,
                    "total_tokens": 0,
                    "cost_usd": 0.0,
                    "estimated": False,
                },
            }

        llm = call_openrouter_extract(
            texto_base,
            api_key,
            model=model,
            filename_hint=nome_arquivo,
            pre_filtrar=pre_filtrar,
            lista_referencia=lista_referencia,
        )
        debug_events.append(f"Biblioteca acionada: requests (modelo IA = {model})")
        debug_events.append(f"Achado IA: {len(llm.get('itens', []))} item(ns) retornado(s)")
        itens = _anotar_fonte_origem(llm.get("itens", []), "ia")
        empresa_det = extrair_razao_social(texto_base)
        debug_highlights.append("RELEVANTE: imagem processada por OCR + IA")

        return {
            "empresa": llm.get("empresa") or empresa_det or os.path.splitext(nome_arquivo)[0],
            "itens": itens,
            "review": review,
            "confianca_estrutural": confianca,
            "fonte_processamento": "ia",
            "texto_truncado": bool(llm.get("texto_truncado")),
            "cnpj": extrair_cnpj(texto_base),
            "telefone": extrair_telefone(texto_base),
            "debug_events": debug_events,
            "debug_highlights": debug_highlights,
            "usage": llm.get("usage", {
                "prompt_tokens": 0,
                "completion_tokens": 0,
                "total_tokens": 0,
                "cost_usd": 0.0,
                "estimated": True,
            }),
        }

    texto_base = extract_text(path)
    llm = call_openrouter_extract(
        texto_base,
        api_key,
        model=model,
        filename_hint=nome_arquivo,
        pre_filtrar=pre_filtrar,
        lista_referencia=lista_referencia,
    )
    debug_events.append(f"Biblioteca acionada: requests (modelo IA = {model})")
    debug_events.append(f"Achado IA: {len(llm.get('itens', []))} item(ns) retornado(s)")
    itens = _anotar_fonte_origem(llm.get("itens", []), "ia")
    empresa_det = extrair_razao_social(texto_base)
    debug_highlights.append("RELEVANTE: arquivo roteado diretamente para IA")

    return {
        "empresa": llm.get("empresa") or empresa_det or os.path.splitext(nome_arquivo)[0],
        "itens": itens,
        "review": review,
        "confianca_estrutural": confianca,
        "fonte_processamento": fonte_global,
        "texto_truncado": bool(llm.get("texto_truncado")),
        "cnpj": extrair_cnpj(texto_base),
        "telefone": extrair_telefone(texto_base),
        "debug_events": debug_events,
        "debug_highlights": debug_highlights,
        "usage": llm.get("usage", {
            "prompt_tokens": 0,
            "completion_tokens": 0,
            "total_tokens": 0,
            "cost_usd": 0.0,
            "estimated": True,
        }),
    }
